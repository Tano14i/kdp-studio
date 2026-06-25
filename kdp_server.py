"""
KDP Studio — Live Trend Backend v2
====================================
Endpoints:
  GET  /health
  GET  /discover          ← ZERO-BIAS: raw global data, no niche pre-selection
  POST /trends            ← niche-guided discovery
  POST /niches
  POST /generate          ← book-type-aware content generation
  POST /generate-all      ← sequential all chapters
  POST /analyze-market    ← competitor intelligence from pasted Amazon data
  POST /title-variants    ← 5 title angles (curiosity/benefit/problem/identity/authority)
  POST /package           ← multilingual KDP package
"""

import os, asyncio, json, random, re as _re
from datetime import datetime
from typing import Optional, List
from fastapi import FastAPI, HTTPException, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import httpx
import anthropic

# ── Config ────────────────────────────────────────────────────
def env(key, default=""):
    return os.environ.get(key, default)

ANTHROPIC_KEY     = env("ANTHROPIC_API_KEY")
REDDIT_USER_AGENT = env("REDDIT_USER_AGENT", "KDPStudio/1.0 (personal use, no auth)")
YOUTUBE_API_KEY   = env("YOUTUBE_API_KEY")
KDP_API_KEY       = env("KDP_API_KEY")   # if set, all POST endpoints require X-API-Key header

# ── SSRF-safe allowed hosts for user-supplied URLs ────────────
_AMAZON_HOSTS = {
    "amazon.com", "www.amazon.com", "amazon.it", "www.amazon.it",
    "amazon.de", "www.amazon.de", "amazon.co.uk", "www.amazon.co.uk",
    "amazon.es", "www.amazon.es", "amazon.fr", "www.amazon.fr",
    "amazon.co.jp", "www.amazon.co.jp", "a.co",
}

def _validate_amazon_url(url: str) -> str:
    """Return url if it points to an allowed Amazon host, else raise."""
    from urllib.parse import urlparse
    import ipaddress
    if not url:
        return url
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError(f"URL scheme not allowed: {parsed.scheme}")
    host = parsed.hostname or ""
    # Block private/loopback IPs
    try:
        ip = ipaddress.ip_address(host)
        if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved:
            raise ValueError("Private/internal IP not allowed")
    except ValueError as e:
        if "not allowed" in str(e):
            raise
    # Whitelist Amazon domains only
    if not any(host == h or host.endswith("." + h) for h in _AMAZON_HOSTS):
        raise ValueError(f"Host not in Amazon whitelist: {host}")
    return url


app = FastAPI(title="KDP Studio API", version="2.0.0", docs_url=None, redoc_url=None)

# CORS: allow Railway domain + localhost dev. Never wildcard in production.
_ALLOWED_ORIGINS = [
    "https://web-production-e6914.up.railway.app",
    "http://localhost:8000",
    "http://127.0.0.1:8000",
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=_ALLOWED_ORIGINS,
    allow_methods=["GET", "POST"],
    allow_headers=["Content-Type", "X-API-Key"],
)


async def require_api_key(request: Request):
    """Dependency: check X-API-Key header if KDP_API_KEY env var is set."""
    if not KDP_API_KEY:
        return  # auth disabled — set KDP_API_KEY in Railway env to enable
    key = request.headers.get("X-API-Key", "")
    if key != KDP_API_KEY:
        raise HTTPException(status_code=401, detail="Invalid or missing API key")

_AUTH = Depends(require_api_key)

claude     = anthropic.Anthropic(api_key=ANTHROPIC_KEY)
claude_async = anthropic.AsyncAnthropic(api_key=ANTHROPIC_KEY)

# ══════════════════════════════════════════════════════════════
# UNIQUENESS ENGINE
# ══════════════════════════════════════════════════════════════

# ── 50+ rotating seed pool per nicchia ───────────────────────
NICHE_SEEDS = {
    "Self-help & Personal growth": [
        "dopamine detox", "nervous system reset", "somatic healing",
        "polyvagal theory", "window of tolerance", "parts work IFS",
        "inner child healing", "shadow work", "ego dissolution",
        "attachment repair", "emotional granularity", "cognitive load",
        "rejection therapy", "learned helplessness", "ego depletion",
        "body doubling ADHD", "delusional confidence", "main character energy",
        "soft life", "romanticize your life", "goblin mode recovery",
        "quiet quitting identity", "rest as resistance", "slow living",
        "digital minimalism", "intentional living", "radical acceptance",
        "self compassion practice", "inner critic work", "identity shifting",
        "neuroplasticity habits", "emotional adulthood", "boundaries healing",
    ],
    "Health & Wellness": [
        "cortisol face", "cycle syncing", "adrenal fatigue recovery",
        "gut brain axis", "seed cycling hormones", "lymphatic drainage",
        "zone 2 cardio", "VO2 max longevity", "cold exposure protocol",
        "mouth taping sleep", "circadian fasting", "autophagy trigger",
        "vagal tone exercises", "breathwork stress", "red light therapy",
        "carnivore healing", "elimination diet", "FODMAP gut health",
        "mineral deficiency symptoms", "hair loss cortisol women",
        "perimenopause symptoms", "testosterone optimization natural",
        "sleep architecture stages", "chronobiology optimization",
    ],
    "Finance & Money": [
        "loud budgeting", "soft saving", "no spend challenge",
        "cash stuffing envelopes", "doom spending anxiety",
        "financial therapy trauma", "money dysmorphia",
        "stealth wealth mindset", "barista FIRE", "coast FIRE",
        "geo arbitrage remote", "house hacking starter",
        "index fund simple path", "bond tent retirement",
        "value averaging strategy", "dividend snowball",
        "frugal hedonic adaptation", "lifestyle inflation trap",
        "financial independence immigrants", "first gen wealth building",
    ],
    "Relationships & Dating": [
        "anxious attachment healing", "fearful avoidant earned security",
        "disorganized attachment", "limerence recovery",
        "situationship clarity", "coercive control recovery",
        "narcissistic abuse healing", "DARVO pattern recognition",
        "love bombing aftermath", "trauma bonding escape",
        "enmeshment untangling", "parentification healing",
        "emotional immaturity parents", "chosen family building",
        "loneliness epidemic men", "male loneliness crisis",
        "female friendship adult", "platonic intimacy",
    ],
    "Productivity & Habits": [
        "body doubling technique", "temptation bundling",
        "implementation intentions", "friction reduction habits",
        "dopamine menu ADHD", "time blindness solutions",
        "ADHD paralysis tools", "executive function scaffolding",
        "ultradian rhythm work", "deep work newport",
        "time blocking template", "weekly review system",
        "second brain building", "progressive summarization",
        "GTD capture trusted", "slow productivity cal newport",
        "anti goals framework", "one thing gary keller",
        "essentialism mckeown", "digital detox productivity",
    ],
    "Spirituality & Mindset": [
        "synchronicity jung meaning", "spiritual bypassing awareness",
        "dark night of soul", "ego death integration",
        "non dual awareness", "present moment power tolle",
        "stoic journaling marcus aurelius", "memento mori practice",
        "amor fati framework", "negative visualization stoic",
        "ikigai finding purpose", "wabi sabi imperfection",
        "mono no aware grief", "ubuntu philosophy community",
        "ubuntu african philosophy", "hygge danish wellbeing",
        "lagom swedish balance", "friluftsliv outdoor mindset",
    ],
    "Fitness & Sports": [
        "hybrid athlete training", "zone 2 fat burning",
        "rucking benefits beginners", "functional fitness aging",
        "strength training women over 40", "perimenopause exercise",
        "protein timing muscle", "creatine women benefits",
        "mobility longevity routine", "fascia release techniques",
        "nervous system recovery training", "HRV optimization",
        "sleep for athletes", "deload week structure",
        "progressive overload beginners", "RPE training scale",
    ],
    "Parenting & Family": [
        "gentle parenting boundaries", "connection before correction",
        "playful parenting cohen", "PACE parenting DDP",
        "therapeutic parenting trauma", "sensory processing kids",
        "twice exceptional 2e parenting", "ADHD girls undiagnosed",
        "autism late diagnosis parenting", "PDA pathological demand",
        "interoception kids body signals", "co regulation adults",
        "reparenting yourself childhood", "adult children estrangement",
        "boomerang generation parenting", "sandwich generation stress",
    ],
    "Business & Entrepreneurship": [
        "solopreneur systems", "one person business model",
        "creator economy monetize", "audience of 1000 fans",
        "niche down riches", "info product validation",
        "cold email renaissance", "linkedin creator B2B",
        "bootstrapped SaaS indie", "micro SaaS niche",
        "productized service agency", "fractional executive model",
        "consulting accelerator positioning", "thought leadership monetize",
        "newsletter monetization beehiiv", "podcast to book funnel",
    ],
    "Neuroscience & Psychology": [
        "interoception body awareness", "neuroception safety",
        "polyvagal ladder states", "window of tolerance trauma",
        "EMDR self administered", "EFT tapping anxiety",
        "DNRS limbic retraining", "somatic experiencing body",
        "sensorimotor psychotherapy", "ACT hexaflex values",
        "schema therapy modes", "DBT skills emotional regulation",
        "mentalizing reflective function", "theory of mind adults",
        "alexithymia emotional awareness", "intolerance uncertainty therapy",
    ],
    "Any niche": [
        "quiet quitting burnout", "soft life boundaries", "delusional confidence",
        "main character energy", "romanticize your life", "dopamine detox",
        "nervous system healing", "situationship clarity", "loud budgeting",
        "cycle syncing hormones", "hybrid athlete", "gentle parenting",
        "solopreneur systems", "ikigai purpose", "cortisol dysregulation",
    ],
    # ── Italian / European market seeds ──────────────────────
    "Italian": [
        "benessere mentale italiano", "stile di vita sano", "mindfulness italia",
        "finanza personale giovani", "investire stipendio", "risparmio intelligente",
        "side hustle italiano", "lavoro da remoto", "freelance italia",
        "relazioni tossiche", "crescita personale", "autostima bassa",
        "come fare soldi online", "creator economy italia", "tiktok monetizzare",
        "alimentazione sana ricette", "dieta mediterranea salute",
        "sport a casa", "yoga principianti", "meditazione guidata",
        "genitoriale consapevole", "figli adolescenti", "coppia crisi",
    ],
    "Dutch": [
        "persoonlijke ontwikkeling", "mentale gezondheid", "burn-out herstel",
        "financieel vrij worden", "sparen jongeren", "beleggen beginners",
        "zelfstandige ondernemer", "thuiswerken productiviteit",
        "relaties verbeteren", "grenzen stellen", "angst overwinnen",
        "gezond eten simpel", "sporten thuis", "mindfulness dagelijks",
    ],
    "German": [
        "persoenliche entwicklung", "mentale gesundheit", "burnout erholung",
        "finanziell frei werden", "geld sparen tipps", "investieren anfaenger",
        "selbststaendig arbeiten", "produktivitaet steigern",
        "beziehungen verbessern", "grenzen setzen", "angst bewaeltigen",
        "gesund ernaehren einfach", "sport zuhause", "achtsamkeit alltag",
    ],
}

def get_language_seeds(keyword: str) -> list:
    """Detect if keyword is non-English and return appropriate seeds."""
    # Simple heuristic: check for common Italian/Dutch/German words
    kw_lower = keyword.lower()
    italian_markers = ['come', 'fare', 'soldi', 'vita', 'salute', 'lavoro', 'amore']
    dutch_markers = ['hoe', 'geld', 'leven', 'werk', 'gezond', 'relatie']
    german_markers = ['wie', 'geld', 'leben', 'arbeit', 'gesund', 'beziehung']
    
    if any(w in kw_lower for w in italian_markers):
        return NICHE_SEEDS.get("Italian", [])
    if any(w in kw_lower for w in dutch_markers):
        return NICHE_SEEDS.get("Dutch", [])
    if any(w in kw_lower for w in german_markers):
        return NICHE_SEEDS.get("German", [])
    return []

# ── 30+ subreddit pool per nicchia ───────────────────────────
NICHE_SUBREDDIT_POOL = {
    "Self-help & Personal growth": [
        "selfimprovement", "getdisciplined", "decidingtobebetter",
        "selfhelp", "productivity", "LifeAdvice", "findapath",
        "DecidingToBeBetter", "Mindfulness", "emotionalintelligence",
        "psychologyofsex", "InternalFamilySystems", "CPTSD",
        "raisedbynarcissists", "adulting", "quarterlifecrisis",
    ],
    "Health & Wellness": [
        "health", "wellness", "longevity", "biohacking", "nutrition",
        "Fasting", "intermittentfasting", "WellnessOver50",
        "PCOS", "Hashimotos", "thyroid", "chronicillness",
        "adhdwomen", "Fibromyalgia", "guthealth",
    ],
    "Finance & Money": [
        "personalfinance", "financialindependence", "frugal",
        "investing", "povertyfinance", "leanfire", "Fire",
        "Bogleheads", "debtfree", "DaveRamsey",
        "FirstTimeHomeBuyer", "churning", "beermoney",
        "digitalnomad", "eupersonalfinance",
    ],
    "Relationships & Dating": [
        "relationship_advice", "dating_advice", "dating",
        "socialskills", "attachment", "BreakUps", "ExNoContact",
        "NarcissisticAbuse", "raisedbynarcissists", "CPTSD",
        "MentalHealthSupport", "lonely", "FriendshipAdvice",
        "adulting", "MenGetsTooLittle",
    ],
    "Productivity & Habits": [
        "productivity", "getdisciplined", "habittracker",
        "nosurf", "digitalminimalism", "ADHD", "adhdwomen",
        "ObsidianMD", "Notion", "gtd", "Zettelkasten",
        "timemanagement", "neurodiversity", "lazyproductivity",
    ],
    "Spirituality & Mindset": [
        "spirituality", "meditation", "mindfulness", "stoicism",
        "awakened", "Jung", "consciousness", "Buddhism",
        "philosophy", "Psychonaut", "Sober", "AlAnon",
        "occult", "witchcraft", "tarot",
    ],
    "Fitness & Sports": [
        "fitness", "loseit", "xxfitness", "running", "bodyweightfitness",
        "weightlifting", "StrongFirst", "overcominggravity",
        "flexibility", "yoga", "swimming", "cycling",
        "intermittentfasting", "carnivore", "veganfitness",
    ],
    "Parenting & Family": [
        "Parenting", "beyondthebump", "mommit", "daddit",
        "raisingkids", "gentleparenting", "SingleParents",
        "autism", "ADHD", "specialneedssupport",
        "EstrangedAdultChild", "AgingParents", "Blended_Families",
    ],
    "Business & Entrepreneurship": [
        "entrepreneur", "smallbusiness", "startups",
        "SideProject", "passive_income", "Entrepreneur",
        "freelance", "digitalnomad", "ecommerce",
        "marketing", "sales", "consulting", "msp",
        "SaaS", "indiehackers",
    ],
    "Neuroscience & Psychology": [
        "psychology", "neuroscience", "cognitivescience",
        "behavioralscience", "therapy", "CPTSD", "BPD",
        "ADHD", "autism", "OCD", "anxiety", "depression",
        "schizoaffective", "bipolar", "mentalhealth",
    ],
    "Any niche": [
        "selfimprovement", "productivity", "wellness",
        "personalfinance", "psychology", "relationship_advice",
        "getdisciplined", "longevity", "entrepreneur", "mindfulness",
    ],
}

# ── Book type templates — specific structure per format ──────
BOOK_TYPE_TEMPLATES = {
    "Guided Journal": {
        "structure": "Each chapter contains: a 1-page intro explaining the theme, 4-6 guided journal prompts with 8-10 lines of writing space each, a reflection box at the end, and an affirmation. No long prose — prompts are the content.",
        "chapter_instruction": "Write the chapter intro (150-200 words), then 5 SPECIFIC journal prompts (each prompt is 1-2 sentences, thought-provoking and personal), then a closing reflection prompt.",
        "length_note": "Chapters are SHORT — prompts, not essays. 400-600 words of actual text per chapter, rest is writing space.",
        "unique_elements": "Include space cues like '[Write here...]' after each prompt. End each chapter with a 1-sentence affirmation in italics.",
    },
    "Workbook": {
        "structure": "Each chapter = one skill or concept. Sections: (1) Concept explanation 300w, (2) Why it matters, (3) Exercise or worksheet with fill-in sections, (4) Action steps checklist, (5) Progress tracker.",
        "chapter_instruction": "Write concept explanation, then create a structured EXERCISE with labeled fields the reader fills in (e.g. 'My current situation:', 'My goal:', 'Action I will take:'). End with a numbered action checklist.",
        "length_note": "Balance: 40% explanation, 60% exercises and worksheets. Make it interactive on the page.",
        "unique_elements": "Use checkbox lists for action steps. Include 'My Notes:' sections. Add a chapter score tracker (1-10 how well they applied the concept).",
    },
    "30-Day Challenge Book": {
        "structure": "Introduction + 30 daily entries. Each day = Day number, Theme, Challenge task (specific action), Reflection prompt, Done checkbox. Days build on each other progressively.",
        "chapter_instruction": "Each DAY entry: Day X header, 1-sentence theme, specific CHALLENGE TASK (what exactly to do today, actionable), 2 reflection prompts, space for notes.",
        "length_note": "Each day entry: 200-300 words max. The book is a daily companion, not an essay collection.",
        "unique_elements": "Start with a Week 1/2/3/4 overview. Include a habit tracker grid. Day 1 should be extremely easy, Day 30 should feel like a milestone.",
    },
    "Planner": {
        "structure": "Goal-setting intro + weekly/monthly spreads. Sections: Annual goals, Monthly intention pages, Weekly layout (Mon-Sun with priorities, tasks, notes), Daily log pages, Reflection pages.",
        "chapter_instruction": "Create a PLANNING SPREAD with labeled sections: Weekly Intention (1 sentence), Top 3 Priorities, Daily task columns (Mon-Sun), Wins of the week, What to improve.",
        "length_note": "Minimal prose — this is functional. Headers, labels, fill-in sections. The reader writes IN it.",
        "unique_elements": "Include a habit tracker, mood tracker, water intake log. Monthly review pages with percentage-completion fields.",
    },
    "Prompt Book": {
        "structure": "Organized by theme/section. Each page = one prompt. Variety of prompt types: writing prompts, thinking prompts, creative prompts, memory prompts. No answers provided.",
        "chapter_instruction": "Write 8-10 VARIED prompts for this chapter theme. Mix: open-ended questions, scenario-based prompts, memory exploration, creative imagination, values clarification. Each prompt on its own conceptual space.",
        "length_note": "Each prompt is 1-3 sentences maximum. The SPACE for the reader to write is the product. Include 8-10 lines of dotted space after each.",
        "unique_elements": "Categorize by difficulty or depth: Surface → Deeper → Core. Include occasional 'Big Question' prompts marked with a star.",
    },
    "Activity Book": {
        "structure": "Chapters by theme. Each chapter mixes: short reading (200w), activity or exercise (specific to the topic), reflection, and a mini-challenge to complete in daily life.",
        "chapter_instruction": "Write a short intro, then design ONE SPECIFIC ACTIVITY (e.g. mapping exercise, letter-writing, body scan, values sort, gratitude list with twist). Make it interactive and time-bound (10-15 min).",
        "length_note": "Activities should have clear instructions: Step 1, Step 2, Step 3. Include a 'You will need:' section if relevant.",
        "unique_elements": "Include estimated time per activity. Add 'Level up' variations for each activity (easier and harder versions).",
    },
    "Self-help Guide": {
        "structure": "Classic non-fiction structure. Each chapter: hook opening, core concept, research or evidence, personal application framework, case study or story, action steps.",
        "chapter_instruction": "Write a full chapter with: compelling opening (hook), main concept explained clearly, 1 research reference or expert insight, practical framework (3-step or similar), concrete action steps, chapter summary.",
        "length_note": "This is prose-heavy. Full paragraphs, narrative flow, 1200-2000 words per chapter.",
        "unique_elements": "End each chapter with 'Key Takeaways' (3 bullets) and 'This Week's Practice' (1 specific action). Include pull-quotes.",
    },
}

def get_book_template(book_type: str) -> dict:
    """Match book type string to template, with fuzzy matching."""
    book_type_lower = book_type.lower()
    for key, template in BOOK_TYPE_TEMPLATES.items():
        if key.lower() in book_type_lower or book_type_lower in key.lower():
            return template
    # Defaults
    if any(w in book_type_lower for w in ["journal", "diary"]):
        return BOOK_TYPE_TEMPLATES["Guided Journal"]
    if any(w in book_type_lower for w in ["work", "exercise"]):
        return BOOK_TYPE_TEMPLATES["Workbook"]
    if any(w in book_type_lower for w in ["plan", "organiz"]):
        return BOOK_TYPE_TEMPLATES["Planner"]
    if any(w in book_type_lower for w in ["challenge", "day"]):
        return BOOK_TYPE_TEMPLATES["30-Day Challenge Book"]
    return BOOK_TYPE_TEMPLATES["Self-help Guide"]

# ── Chapter opening styles pool ───────────────────────────────
OPENING_STYLES = [
    "Start with a vivid real-world scenario or micro-story (2-3 sentences) that drops the reader into the feeling",
    "Open with a provocative question that challenges a common assumption the reader holds",
    "Begin with a surprising statistic or research finding that reframes the chapter topic",
    "Open with a short quote from a non-English cultural tradition (philosophy, literature, folk wisdom)",
    "Start with a brief personal confession or vulnerability that immediately builds trust",
    "Open with a paradox or contradiction that the chapter will resolve",
    "Begin with a sensory description — a smell, sound, texture — that anchors the emotional tone",
    "Start with a bold, counter-intuitive claim that the reader will want to disprove or explore",
    "Open with a dialogue snippet or overheard conversation that captures the chapter's tension",
    "Begin with a very short poem, haiku, or lyrical fragment that sets the emotional key",
]

# ── Writing tone pool ─────────────────────────────────────────
TONE_DESCRIPTORS = [
    "warm and conversational, like a trusted friend who happens to be an expert",
    "precise and direct, no fluff — every sentence earns its place",
    "lyrical and reflective, with space for the reader to breathe between ideas",
    "academically grounded but emotionally accessible — cite concepts without jargon",
    "bold and challenging, pushing the reader past their comfort zone with compassion",
    "gentle and trauma-informed, never rushing, always meeting the reader where they are",
    "intellectually curious, exploring ideas from multiple cultural angles",
    "pragmatic and systems-oriented, always asking: what does the reader do next?",
]

# ── Uniqueness helpers ────────────────────────────────────────
def now_stamp() -> str:
    n = datetime.now()
    return f"{n.strftime('%A %d %B %Y')}, {n.strftime('%H:%M')} CET — week {n.isocalendar()[1]} of {n.year}"

# Timeframe → Reddit + Google Trends params
TIMEFRAME_MAP = {
    "day":      {"reddit_t": "day",   "gtrends": "now 1-d",    "label": "Oggi",          "strategy": "Pubblica entro 1-2 settimane o il momento passa"},
    "week":     {"reddit_t": "week",  "gtrends": "now 7-d",    "label": "Settimana",     "strategy": "Finestra ottimale KDP — caldo ma non ancora saturo"},
    "month":    {"reddit_t": "month", "gtrends": "today 1-m",  "label": "Mese",          "strategy": "Nicchia in crescita — meno urgenza, audience più definita"},
    "3months":  {"reddit_t": "month", "gtrends": "today 3-m",  "label": "3 Mesi",        "strategy": "Trend in consolidamento — cerca angoli non ancora coperti"},
    "year":     {"reddit_t": "year",  "gtrends": "today 12-m", "label": "Anno",          "strategy": "Nicchia matura — migliora libri esistenti o trova sub-nicchie"},
}

def get_timeframe(tf: str) -> dict:
    return TIMEFRAME_MAP.get(tf, TIMEFRAME_MAP["week"])

def pick_seeds(niche: str, keyword: str = "", n: int = 4) -> list[str]:
    pool = []
    # First check if keyword suggests a non-English language
    if keyword:
        lang_seeds = get_language_seeds(keyword)
        if lang_seeds:
            pool = lang_seeds
    # Then try niche-based pool
    if not pool:
        for k, v in NICHE_SEEDS.items():
            if k.lower() in niche.lower() or niche.lower() in k.lower():
                pool = v
                break
    if not pool:
        pool = NICHE_SEEDS["Any niche"]
    seeds = random.sample(pool, min(n, len(pool)))
    if keyword and keyword not in seeds:
        seeds[0] = keyword  # always include user keyword
    return seeds

def pick_subreddits(niche: str, n: int = 4) -> list[str]:
    pool = []
    for k, v in NICHE_SUBREDDIT_POOL.items():
        if k.lower() in niche.lower() or niche.lower() in k.lower():
            pool = v
            break
    if not pool:
        pool = NICHE_SUBREDDIT_POOL["Any niche"]
    return random.sample(pool, min(n, len(pool)))

def pick_opening_style() -> str:
    return random.choice(OPENING_STYLES)

def pick_tone() -> str:
    return random.choice(TONE_DESCRIPTORS)

# ══════════════════════════════════════════════════════════════
# PYDANTIC MODELS
# ══════════════════════════════════════════════════════════════
MARKET_LANG_CONFIG = {
    "English":    {"amazon": "Amazon.com", "subreddits": ["books","selfhelp","personalfinance","productivity","Fitness","mentalhealth","relationships","Parenting","Entrepreneur"]},
    "Italian":    {"amazon": "Amazon.it",  "subreddits": ["italy","italiani","crescitapersonale","ItalyInformatica","psicologia"]},
    "Spanish":    {"amazon": "Amazon.es / Amazon.com.mx", "subreddits": ["es","mexico","argentina","colombia","emprendimiento","psicologia"]},
    "German":     {"amazon": "Amazon.de",  "subreddits": ["de","Austria","finanzen","Persoenlichkeitsentwicklung","Switzerland"]},
    "French":     {"amazon": "Amazon.fr",  "subreddits": ["france","francophonie","Quebec","developpementpersonnel"]},
    "Portuguese": {"amazon": "Amazon.com.br", "subreddits": ["brasil","portugal","empreendedorismo","desabafos"]},
}

class TrendRequest(BaseModel):
    platforms: list[str]
    niche: str
    keyword: Optional[str] = ""
    timeframe: Optional[str] = "week"   # day|week|month|3months|year
    market_language: Optional[str] = "English"

class NicheRequest(BaseModel):
    platforms: list[str]
    keyword: Optional[str] = ""
    timeframe: Optional[str] = "week"

class PositioningRequest(BaseModel):
    author_identity: str       # chi sei come autore/publisher, come vuoi essere percepito
    ideal_reader: str           # psicografia del lettore ideale (non demografica)
    transformation: str         # trasformazione specifica che prometti
    competitors: Optional[str] = ""      # competitor diretti e differenziazione
    unfair_advantage: Optional[str] = "" # cosa sai/hai vissuto che altri non hanno

class GenerateRequest(BaseModel):
    trend_name: str
    book_title: str
    book_subtitle: str
    book_type: str
    audience: str
    tab: str
    chapter_num: Optional[int] = 1
    outline: Optional[str] = ""
    # New voice params
    tone: Optional[str] = ""
    language: Optional[str] = "English"
    cultural_inspiration: Optional[str] = ""
    chapter_length: Optional[str] = "medium"   # short|medium|long
    reader_persona: Optional[str] = ""
    custom_instructions: Optional[str] = ""
    amazon_keywords: Optional[List[str]] = []
    niche_analysis: Optional[dict] = None   # best_angle, opportunities, keyword_gems, etc.

class AllChaptersRequest(BaseModel):
    trend_name: str
    book_title: str
    book_subtitle: str
    book_type: str
    audience: str
    outline: str
    tone: Optional[str] = ""
    language: Optional[str] = "English"
    cultural_inspiration: Optional[str] = ""
    chapter_length: Optional[str] = "medium"
    reader_persona: Optional[str] = ""
    custom_instructions: Optional[str] = ""
    amazon_keywords: Optional[List[str]] = []
    niche_analysis: Optional[dict] = None
    start_from: Optional[int] = 1  # resume support: skip chapters before this number

class PackageRequest(BaseModel):
    trend_name: str
    trend_platform: str
    book_title: str
    book_subtitle: str
    book_type: str
    audience: str
    tone: Optional[str] = ""
    reader_persona: Optional[str] = ""
    language: Optional[str] = "English"
    custom_instructions: Optional[str] = ""

# ══════════════════════════════════════════════════════════════
LANG_CODE_MAP = {
    "English": "en", "Italian": "it", "Spanish": "es",
    "German": "de", "French": "fr", "Portuguese": "pt",
}

# ══════════════════════════════════════════════════════════════
# GOOGLE / YOUTUBE AUTOCOMPLETE HELPERS
# ══════════════════════════════════════════════════════════════
async def fetch_google_autocomplete(query: str, lang_code: str = "en") -> list[str]:
    try:
        async with httpx.AsyncClient(timeout=8) as client:
            r = await client.get(
                "https://suggestqueries.google.com/complete/search",
                params={"client": "firefox", "q": query, "hl": lang_code},
                headers={"Accept-Language": lang_code}
            )
            if r.status_code != 200:
                return []
            data = r.json()
            return [s for s in (data[1] if len(data) > 1 else []) if s.lower() != query.lower()][:10]
    except Exception as e:
        print(f"[Google Autocomplete] {query}: {e}")
        return []

async def fetch_youtube_autocomplete(query: str, lang_code: str = "en") -> list[str]:
    try:
        async with httpx.AsyncClient(timeout=8) as client:
            r = await client.get(
                "https://suggestqueries.google.com/complete/search",
                params={"client": "firefox", "ds": "yt", "q": query, "hl": lang_code},
                headers={"Accept-Language": lang_code}
            )
            if r.status_code != 200:
                return []
            data = r.json()
            return [s for s in (data[1] if len(data) > 1 else []) if s.lower() != query.lower()][:10]
    except Exception as e:
        print(f"[YouTube Autocomplete] {query}: {e}")
        return []

async def fetch_multi_autocomplete(niche: str, lang_code: str = "en") -> dict:
    """Batch Google + YouTube autocomplete across multiple seed queries."""
    kw = niche.strip()
    seeds_google = [kw, f"{kw} book", f"{kw} guide", f"how to {kw}", f"best {kw}"]
    seeds_yt = [kw, f"how to {kw}", f"{kw} tips"]
    tasks = [fetch_google_autocomplete(s, lang_code) for s in seeds_google]
    tasks += [fetch_youtube_autocomplete(s, lang_code) for s in seeds_yt]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    google_seen, google_all = set(), []
    for r in results[:len(seeds_google)]:
        if isinstance(r, list):
            for s in r:
                if s not in google_seen:
                    google_seen.add(s); google_all.append(s)

    yt_seen, yt_all = set(), []
    for r in results[len(seeds_google):]:
        if isinstance(r, list):
            for s in r:
                if s not in yt_seen:
                    yt_seen.add(s); yt_all.append(s)

    current_year = datetime.now().year
    def _stale(kw: str) -> bool:
        years = [int(m) for m in _re.findall(r'\b(20\d{2})\b', kw)]
        return any(y <= current_year - 2 for y in years)

    google_all = [k for k in google_all if not _stale(k)]
    yt_all = [k for k in yt_all if not _stale(k)]
    return {"google": google_all[:25], "youtube": yt_all[:15]}

# ══════════════════════════════════════════════════════════════
# REDDIT HELPER
# ══════════════════════════════════════════════════════════════
async def fetch_reddit_posts(niche: str, keyword: str = "", timeframe: str = "week", force_subreddits: list = None) -> list[dict]:
    if force_subreddits:
        subreddits = random.sample(force_subreddits, min(4, len(force_subreddits)))
    else:
        subreddits = pick_subreddits(niche, n=4)
    tf = get_timeframe(timeframe)
    reddit_t = tf["reddit_t"]
    print(f"[Reddit] Subreddits: {subreddits} | timeframe: {reddit_t}")
    posts = []

    headers = {
        "User-Agent": REDDIT_USER_AGENT,
        "Accept": "application/json",
    }

    async with httpx.AsyncClient(timeout=12, headers=headers, follow_redirects=True) as client:
        async def fetch_sub(sub: str, sort: str) -> list[dict]:
            try:
                url = f"https://www.reddit.com/r/{sub}/{sort}.json"
                params = {"limit": 25, "raw_json": 1}
                if sort == "top":
                    params["t"] = reddit_t
                r = await client.get(url, params=params)
                if r.status_code == 429:
                    await asyncio.sleep(2)
                    r = await client.get(url, params=params)
                if r.status_code != 200:
                    return []
                result = []
                for item in r.json().get("data", {}).get("children", []):
                    p = item.get("data", {})
                    if p.get("stickied") or p.get("is_meta"):
                        continue
                    title = p.get("title", "")
                    if keyword and keyword.lower() not in (title + p.get("selftext","")).lower():
                        continue
                    result.append({
                        "subreddit": p.get("subreddit",""),
                        "title": title,
                        "score": p.get("score", 0),
                        "num_comments": p.get("num_comments", 0),
                        "upvote_ratio": p.get("upvote_ratio", 0),
                        "url": "https://reddit.com" + p.get("permalink",""),
                        "flair": p.get("link_flair_text",""),
                    })
                return result
            except Exception as e:
                print(f"[Reddit] r/{sub}/{sort}: {e}")
                return []

        tasks = [fetch_sub(sub, sort) for sub in subreddits for sort in ["top","hot"]]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        for r in results:
            if isinstance(r, list):
                posts.extend(r)

    seen, unique = set(), []
    for p in sorted(posts, key=lambda x: x["score"], reverse=True):
        if p["title"] not in seen:
            seen.add(p["title"])
            unique.append(p)
    print(f"[Reddit] {len(unique)} unique posts fetched")
    return unique[:25]

# ══════════════════════════════════════════════════════════════
# GOOGLE TRENDS HELPER
# ══════════════════════════════════════════════════════════════
async def fetch_google_trends(niche: str, keyword: str = "", timeframe: str = "week") -> dict:
    """Fast pytrends call for main discovery flow. Fails quickly if Google blocks —
    N/A is acceptable there. Use fetch_google_trends_apify() for validate-niche."""
    _EMPTY = {"terms": [], "avg_interest": {}, "rising_queries": {}}
    try:
        from pytrends.request import TrendReq
        terms = pick_seeds(niche, keyword, n=4)
        tf = get_timeframe(timeframe)
        gtrends_tf = tf["gtrends"]

        def _fetch():
            pt = TrendReq(hl='en-US', tz=360, timeout=(4, 10))
            pt.build_payload(terms, timeframe=gtrends_tf, geo='')
            interest = pt.interest_over_time()
            rising = {}
            try:
                related = pt.related_queries()
                for t in terms:
                    if t in related and related[t].get("rising") is not None:
                        df = related[t]["rising"]
                        rising[t] = df.head(5).to_dict('records') if not df.empty else []
            except Exception:
                pass
            avg = {}
            if not interest.empty:
                for t in terms:
                    if t in interest.columns:
                        avg[t] = int(interest[t].mean())
            return {"terms": terms, "avg_interest": avg, "rising_queries": rising}

        loop = asyncio.get_event_loop()
        result = await asyncio.wait_for(loop.run_in_executor(None, _fetch), timeout=12)
        return result
    except asyncio.TimeoutError:
        print("[GTrends] Timeout — Google rate-limiting this IP")
        return _EMPTY
    except Exception as e:
        print(f"[GTrends] Unavailable: {type(e).__name__}")
        return _EMPTY


def _parse_traffic(s: str) -> int:
    """Convert '100+', '10K+', '1M+' traffic strings to integers."""
    s = str(s).upper().replace("+", "").replace(",", "").strip()
    mult = 1
    if s.endswith("K"):
        s, mult = s[:-1], 1_000
    elif s.endswith("M"):
        s, mult = s[:-1], 1_000_000
    try:
        return int(float(s) * mult)
    except (ValueError, TypeError):
        return 0


async def fetch_google_trends_apify(niche: str, keyword: str = "",
                                     market_language: str = "English") -> dict:
    """Google Trends data via Apify. The automation-lab actor returns trending
    topics (type='trending') rather than interest-over-time, so we:
    1. Check if any of our seed terms appear in the trending list → interest score
    2. Return all trending keywords for the geo as rising_queries (genuinely useful)
    3. Preserve interest-over-time parsing for if/when we switch actors."""
    _EMPTY = {"terms": [], "avg_interest": {}, "rising_queries": {}, "trend_direction": None}
    if not APIFY_TOKEN:
        return _EMPTY

    terms = pick_seeds(niche, keyword, n=3)
    geo = MARKET_GEO_MAP.get(market_language, "US")

    try:
        items = await run_actor(
            "automation-lab/google-trends-scraper",
            {"geo": geo, "timeRange": "today 3-m"},
            timeout_sec=90,
        )
    except Exception as e:
        print(f"[GTrends/Apify] Error: {e}")
        return _EMPTY

    if not items:
        return _EMPTY

    avg_interest: dict[str, int] = {}
    rising_queries: dict[str, list] = {}
    trend_direction: dict[str, str] = {}
    trending_pool: list[dict] = []   # all trending items for this geo

    for item in items:
        if not isinstance(item, dict):
            continue

        item_type = item.get("type", "")

        # ── Trending-topics format (automation-lab actor default output) ──
        if item_type == "trending":
            kw = item.get("keyword", "")
            traffic = _parse_traffic(item.get("traffic", "0"))
            if kw:
                trending_pool.append({"query": kw, "value": str(traffic)})
                # If one of our seed terms matches a trending topic → score it
                for term in terms:
                    if term.lower() in kw.lower() or kw.lower() in term.lower():
                        # Scale: 100 traffic → ~30, 1K → ~55, 10K → ~75, 1M → ~100
                        import math
                        score = min(100, max(10, int(30 + 10 * math.log10(max(traffic, 1)))))
                        avg_interest[term] = max(avg_interest.get(term, 0), score)
            continue

        # ── Interest-over-time format (future actor or different mode) ──
        kw = (item.get("keyword") or item.get("term") or
              item.get("searchTerm") or item.get("query") or "")
        if not kw:
            continue

        timeline = (item.get("timeline") or item.get("interestOverTime") or
                    item.get("interest_over_time") or item.get("data") or [])
        if timeline:
            vals = []
            for p in timeline:
                if not isinstance(p, dict):
                    continue
                try:
                    vals.append(int(str(p.get("value", 0)).replace("<1", "0")))
                except (ValueError, TypeError):
                    pass
            if vals:
                avg_interest[kw] = round(sum(vals) / len(vals))
                mid = len(vals) // 2
                fh = sum(vals[:mid]) / max(mid, 1)
                sh = sum(vals[mid:]) / max(len(vals) - mid, 1)
                trend_direction[kw] = "up" if sh > fh * 1.15 else ("down" if sh < fh * 0.85 else "stable")

        rq = item.get("relatedQueries") or item.get("related_queries") or {}
        rising = rq.get("rising") or []
        if rising:
            rising_queries[kw] = [
                {"query": r.get("query", r) if isinstance(r, dict) else str(r),
                 "value": r.get("value", "") if isinstance(r, dict) else ""}
                for r in rising[:5]
            ]

    # Trending pool → rising_queries (top items by traffic, useful for the user)
    if trending_pool and not rising_queries:
        try:
            sorted_pool = sorted(trending_pool, key=lambda x: -int(x["value"]) if x["value"].isdigit() else 0)
        except Exception:
            sorted_pool = trending_pool
        primary = terms[0] if terms else niche
        rising_queries[primary] = sorted_pool[:8]

    print(f"[GTrends/Apify] geo={geo} trending={len(trending_pool)} matched={list(avg_interest.keys())}")

    return {
        "terms": terms,
        "avg_interest": avg_interest,
        "rising_queries": rising_queries,
        "trend_direction": trend_direction,
        "geo": geo,
    }

# ══════════════════════════════════════════════════════════════
# CLAUDE HELPER
# ══════════════════════════════════════════════════════════════
async def call_claude(prompt: str, max_tokens: int = 4000, allow_truncated: bool = False) -> str:
    """Call Claude with a 90s timeout and 1 retry on transient errors (overload/rate-limit)."""
    _RETRYABLE = {529, 529, 500, 503}  # HTTP-equivalent stop reasons / status codes
    for attempt in range(2):
        try:
            msg = await asyncio.wait_for(
                claude_async.messages.create(
                    model="claude-sonnet-4-6",
                    max_tokens=max_tokens,
                    messages=[{"role": "user", "content": prompt}]
                ),
                timeout=90,
            )
            if msg.stop_reason == "max_tokens" and not allow_truncated:
                raise ValueError("Risposta troncata — usa Capitolo Singolo o riduci la lunghezza.")
            return msg.content[0].text
        except asyncio.TimeoutError:
            if attempt == 0:
                await asyncio.sleep(3)
                continue
            raise HTTPException(status_code=504, detail="Claude timeout — riprova tra qualche secondo.")
        except Exception as e:
            code = getattr(e, 'status_code', None)
            if attempt == 0 and code in (429, 529, 500, 503):
                await asyncio.sleep(5)
                continue
            raise

def parse_json_safe(text: str) -> dict:
    import re, logging as _log
    text = text.replace('\u2018',"'").replace('\u2019',"'")
    text = text.replace('\u201c','"').replace('\u201d','"')
    text = text.replace('\u2014','-').replace('\u2013','-')
    m = re.search(r'\{[\s\S]*\}', text)
    if not m:
        raise ValueError("Nessun JSON nella risposta Claude")
    j = m.group(0)
    try:
        return json.loads(j)
    except json.JSONDecodeError as e:
        # "Extra data": valid JSON followed by trailing text \u2014 truncate at e.pos
        if e.pos and e.pos > 0:
            try:
                result = json.loads(j[:e.pos])
                _log.warning("parse_json_safe: truncated trailing data at pos %d", e.pos)
                return result
            except json.JSONDecodeError:
                pass
        _log.warning("parse_json_safe: applying regex repair (trailing comma / control chars)")
        j = re.sub(r',\s*([\]}])', r'\1', j)
        j = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', j)
        return json.loads(j)

def build_voice_ctx(tone="", language="English", cultural_inspiration="",
                    chapter_length="medium", reader_persona="", custom_instructions="") -> str:
    """Build a voice/style context block to inject into every content prompt."""
    length_map = {"short": "800-1000", "medium": "1200-1600", "long": "2000-2500"}
    word_count = length_map.get(chapter_length, "1200-1600")
    opening = pick_opening_style()
    writing_tone = tone if tone else pick_tone()
    lines = [
        f"WRITING VOICE & STYLE (follow precisely):",
        f"- Tone: {writing_tone}",
        f"- Output language: {language}",
        f"- Target word count: {word_count} words",
        f"- Chapter opening style: {opening}",
    ]
    if cultural_inspiration:
        lines.append(f"- Cultural lens: draw references, metaphors, or philosophy from {cultural_inspiration} tradition")
    if reader_persona:
        lines.append(f"- Write as if speaking directly to: {reader_persona}")
    if custom_instructions:
        lines.append(f"\nCUSTOM RESTRICTIONS (MANDATORY — apply to every sentence):\n{custom_instructions}")
    lines.append(f"- Uniqueness seed: {now_stamp()} — use this to make this version distinct from any previous generation")
    lines.append(f"- IMPORTANT: Do NOT use generic self-help clichés. Make every sentence specific and surprising.")
    return "\n".join(lines)

# ══════════════════════════════════════════════════════════════
# ROUTES
# ══════════════════════════════════════════════════════════════
# ── Serve frontend HTML (so Railway hosts everything) ────────
@app.get("/")
async def serve_frontend():
    from fastapi.responses import FileResponse
    import os
    # Look for the frontend HTML in same dir as server
    base_dir = os.path.dirname(os.path.abspath(__file__))
    for name in ("kdp-trend-hunter.html", "index.html"):
        html_path = os.path.join(base_dir, name)
        if os.path.exists(html_path):
            return FileResponse(html_path, media_type="text/html")
    return {"message": "KDP Studio API running. Frontend not found — place kdp-trend-hunter.html in same folder."}


@app.get("/health")
async def health():
    # Return immediately — Reddit status checked async in background
    # This prevents the health check from timing out while waiting for Reddit
    return {
        "status": "ok",
        "reddit": True,   # assumed ok — verified at request time
        "reddit_mode": "public_json",
        "claude": bool(ANTHROPIC_KEY),
        "timestamp": datetime.utcnow().isoformat()
    }


@app.get("/health/full")
async def health_full():
    # Full health check including Reddit ping (slower, use only when needed)
    reddit_ok = False
    try:
        async with httpx.AsyncClient(timeout=5, headers={"User-Agent": REDDIT_USER_AGENT}) as c:
            r = await c.get("https://www.reddit.com/r/selfimprovement/hot.json?limit=1&raw_json=1")
            reddit_ok = r.status_code == 200
    except Exception:
        pass
    tiktok_ok = False
    try:
        async with httpx.AsyncClient(timeout=5) as c:
            r = await c.get("https://ads.tiktok.com/creative_radar_api/v1/popular_trend/hashtag/list", params={
                "page": 1, "limit": 1, "period": 7, "country_code": "US", "sort_by": "popular",
            }, headers={"User-Agent": "Mozilla/5.0", "Accept": "application/json"})
            tiktok_ok = r.status_code == 200
    except Exception:
        pass
    amazon_autocomplete_ok = False
    try:
        async with httpx.AsyncClient(timeout=5, headers={**AMAZON_HEADERS, "Accept": "application/json"}) as c:
            r = await c.get("https://completion.amazon.com/api/2017/suggestions", params={
                "limit": 1, "prefix": "journal", "suggestion-type": "KEYWORD",
                "page-type": "Search", "alias": "stripbooks", "mid": "ATVPDKIKX0DER",
            })
            amazon_autocomplete_ok = r.status_code == 200
    except Exception:
        pass
    return {
        "status": "ok",
        "reddit": reddit_ok,
        "reddit_mode": "public_json",
        "claude": bool(ANTHROPIC_KEY),
        "youtube": bool(YOUTUBE_API_KEY),
        "tiktok": tiktok_ok,
        "amazon_autocomplete": amazon_autocomplete_ok,
        "timestamp": datetime.utcnow().isoformat()
    }


# ══════════════════════════════════════════════════════════════
# ZERO-BIAS DISCOVERY ENGINE
# ══════════════════════════════════════════════════════════════

# Subreddits that reflect general human discourse — no niche bias
DISCOVERY_SUBREDDITS = [
    # Human experience & emotions
    "TrueOffMyChest", "offmychest", "self", "AITA", "confessions",
    "LifeAdvice", "NoStupidQuestions", "AskReddit", "Showerthoughts",
    # Behavior & change
    "ChangeMyView", "unpopularopinion", "raisedbynarcissists",
    "relationship_advice", "socialskills", "lonely",
    # Books & learning
    "books", "writing", "suggestmeabook", "NonFictionBookClub",
    # Broad wellness signals
    "mentalhealth", "anxiety", "depression", "selfimprovement",
    "findapath", "careerguidance", "adulting",
]

# Google Trends country codes to rotate (no topic seed)
DISCOVERY_COUNTRIES = ["US", "GB", "AU", "CA", "NL", "DE", "IT"]

async def fetch_reddit_global(limit_per_sub: int = 15) -> list[dict]:
    """
    Fetch top posts from r/all + a random mix of broad subreddits.
    No niche filtering whatsoever.
    """
    # Pick 5 random broad subreddits + always include r/all
    subs = ["all"] + random.sample(DISCOVERY_SUBREDDITS, min(5, len(DISCOVERY_SUBREDDITS)))
    print(f"[Discovery] Subreddits: {subs}")
    posts = []
    headers = {"User-Agent": REDDIT_USER_AGENT, "Accept": "application/json"}

    async with httpx.AsyncClient(timeout=12, headers=headers, follow_redirects=True) as client:
        async def fetch_one(sub: str, sort: str) -> list[dict]:
            try:
                url = f"https://www.reddit.com/r/{sub}/{sort}.json"
                params = {"limit": limit_per_sub, "raw_json": 1}
                if sort == "top":
                    params["t"] = "day"  # always today for zero-bias discovery
                r = await client.get(url, params=params)
                if r.status_code == 429:
                    await asyncio.sleep(2)
                    r = await client.get(url, params=params)
                if r.status_code != 200:
                    return []
                result = []
                for item in r.json().get("data", {}).get("children", []):
                    p = item.get("data", {})
                    if p.get("stickied") or p.get("is_meta"):
                        continue
                    # Skip purely entertainment/meme posts
                    if p.get("subreddit","").lower() in ["funny","memes","pics","gifs","videos"]:
                        continue
                    result.append({
                        "subreddit": p.get("subreddit",""),
                        "title": p.get("title",""),
                        "score": p.get("score", 0),
                        "num_comments": p.get("num_comments", 0),
                        "upvote_ratio": p.get("upvote_ratio", 0),
                        "url": "https://reddit.com" + p.get("permalink",""),
                        "selftext_snippet": p.get("selftext","")[:200],
                    })
                return result
            except Exception as e:
                print(f"[Discovery] r/{sub}/{sort}: {e}")
                return []

        tasks = [fetch_one(sub, sort) for sub in subs for sort in ["top","rising"]]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        for r in results:
            if isinstance(r, list):
                posts.extend(r)

    # Deduplicate, sort by score
    seen, unique = set(), []
    for p in sorted(posts, key=lambda x: x["score"], reverse=True):
        if p["title"] not in seen:
            seen.add(p["title"])
            unique.append(p)

    print(f"[Discovery] {len(unique)} unique posts from global fetch")
    return unique[:40]


async def fetch_google_trending_now() -> dict:
    """
    Fetch what is trending on Google RIGHT NOW — no seed, no topic bias.
    Uses pytrends trending_searches and realtime_trending_searches.
    """
    try:
        from pytrends.request import TrendReq
        country = random.choice(DISCOVERY_COUNTRIES)
        print(f"[Discovery] Google trending in: {country}")

        def _fetch():
            pt = TrendReq(hl="en-US", tz=360, timeout=(10,25))
            results = {}

            # Daily trending searches — top ~20 queries of today
            try:
                df = pt.trending_searches(pn="united_states")
                results["daily_trending"] = df[0].tolist()[:20] if not df.empty else []
            except Exception as e:
                print(f"[GTrends Daily] {e}")
                results["daily_trending"] = []

            # Realtime trending — what is exploding RIGHT NOW
            try:
                rt = pt.realtime_trending_searches(pn="US")
                if rt is not None and not rt.empty:
                    titles = []
                    for col in ["title", "entityNames", "query"]:
                        if col in rt.columns:
                            titles = rt[col].dropna().tolist()[:15]
                            break
                    results["realtime"] = [str(t) for t in titles]
                else:
                    results["realtime"] = []
            except Exception as e:
                print(f"[GTrends Realtime] {e}")
                results["realtime"] = []

            return results

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _fetch)

    except Exception as e:
        print(f"[GoogleTrending] {e}")
        return {"daily_trending": [], "realtime": [], "error": str(e)}


async def fetch_youtube_trending(limit: int = 25) -> list[dict]:
    """
    Fetch YouTube's official Trending feed (chart=mostPopular) — no search
    query, no niche bias. Requires YOUTUBE_API_KEY (free quota via
    Google Cloud Console). Returns [] if not configured or on error.
    """
    if not YOUTUBE_API_KEY:
        return []
    try:
        country = random.choice(["US", "GB", "AU", "CA"])
        async with httpx.AsyncClient(timeout=10) as client:
            r = await client.get("https://www.googleapis.com/youtube/v3/videos", params={
                "part": "snippet,statistics",
                "chart": "mostPopular",
                "maxResults": limit,
                "regionCode": country,
                "key": YOUTUBE_API_KEY,
            })
            if r.status_code != 200:
                print(f"[YouTube] {r.status_code}: {r.text[:200]}")
                return []
            items = []
            for v in r.json().get("items", []):
                sn, st = v.get("snippet", {}), v.get("statistics", {})
                items.append({
                    "title": sn.get("title", ""),
                    "category_id": sn.get("categoryId", ""),
                    "channel": sn.get("channelTitle", ""),
                    "views": int(st.get("viewCount", 0)),
                    "region": country,
                })
            return items
    except Exception as e:
        print(f"[YouTube] {e}")
        return []


AMAZON_HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Accept-Language": "en-US,en;q=0.9",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
}

async def fetch_amazon_book_info(url: str) -> dict:
    """Follow redirect and extract title/author/description from an Amazon product page."""
    if not url:
        return {}
    try:
        url = _validate_amazon_url(url)
    except ValueError as e:
        return {"error": str(e)}
    try:
        async with httpx.AsyncClient(timeout=15, follow_redirects=True) as client:
            r = await client.get(url, headers=AMAZON_HEADERS)
        html = r.text

        import re, json as _json

        # Try JSON-LD first (most reliable)
        for m in re.finditer(r'<script[^>]*type="application/ld\+json"[^>]*>(.*?)</script>', html, re.DOTALL):
            try:
                obj = _json.loads(m.group(1))
                if isinstance(obj, list):
                    obj = obj[0]
                if obj.get("@type") in ("Book", "Product"):
                    info: dict = {}
                    if obj.get("name"):
                        info["title"] = obj["name"]
                    if obj.get("description"):
                        info["description"] = re.sub(r'\s+', ' ', re.sub(r'<[^>]+>', '', obj["description"])).strip()[:600]
                    if obj.get("author"):
                        a = obj["author"]
                        info["author"] = a.get("name", "") if isinstance(a, dict) else str(a)
                    if info.get("title"):
                        return info
            except Exception:
                pass

        # Fallback: regex scrape
        info = {}
        t = re.search(r'id="productTitle"[^>]*>\s*(.*?)\s*</span>', html, re.DOTALL)
        if t:
            info["title"] = re.sub(r'\s+', ' ', t.group(1)).strip()
        d = re.search(r'id="bookDescription_feature_div".*?<span[^>]*>(.*?)</span>', html, re.DOTALL)
        if not d:
            d = re.search(r'id="productDescription".*?<p[^>]*>(.*?)</p>', html, re.DOTALL)
        if d:
            info["description"] = re.sub(r'\s+', ' ', re.sub(r'<[^>]+>', '', d.group(1))).strip()[:600]
        return info
    except Exception:
        return {}


async def fetch_amazon_book_count(query: str) -> dict:
    """
    Gap analysis: how many books on Amazon already match this niche.
    Scrapes the Amazon Books search results count. Returns count=None
    if blocked or unparseable — caller must handle gracefully.
    """
    import re
    try:
        async with httpx.AsyncClient(timeout=10, headers=AMAZON_HEADERS, follow_redirects=True) as client:
            r = await client.get("https://www.amazon.com/s", params={"k": query, "i": "stripbooks"})
            if r.status_code != 200:
                return {"query": query, "count": None}
            html = r.text
            m = re.search(r'of\s+(?:over\s+)?([\d,]+)\s+results', html, re.IGNORECASE)
            if not m:
                m = re.search(r'([\d,]+)\s+results?\s+for', html, re.IGNORECASE)
            if m:
                return {"query": query, "count": int(m.group(1).replace(",", ""))}
            return {"query": query, "count": None}
    except Exception as e:
        print(f"[Amazon Gap] {query}: {e}")
        return {"query": query, "count": None}


def classify_amazon_saturation(count: Optional[int]) -> tuple[str, str]:
    if count is None:
        return "unknown", "Dati Amazon non disponibili — verifica manualmente su Amazon"
    if count < 200:
        return "opportunity", f"Solo {count} libri su Amazon in questa nicchia — bassa concorrenza, buona opportunita"
    elif count < 1500:
        return "moderate", f"{count} libri su Amazon — concorrenza moderata, serve un angolo specifico"
    else:
        return "saturated", f"{count}+ libri su Amazon — nicchia satura, cerca una sotto-nicchia o angolo molto diverso"


async def fetch_amazon_autocomplete(query: str) -> list[str]:
    """
    Real demand signal: what Amazon's search-bar autocomplete suggests for this
    niche query — reflects actual searches typed by shoppers/readers in the
    Books store. Undocumented endpoint, degrades to [] on any failure.
    """
    try:
        headers = {**AMAZON_HEADERS, "Accept": "application/json"}
        async with httpx.AsyncClient(timeout=8, headers=headers) as client:
            r = await client.get("https://completion.amazon.com/api/2017/suggestions", params={
                "limit": 10,
                "prefix": query,
                "suggestion-type": "KEYWORD",
                "page-type": "Search",
                "alias": "stripbooks",
                "site-variant": "desktop",
                "version": "3",
                "event": "onKeyPress",
                "wc": "",
                "lop": "en_US",
                "mid": "ATVPDKIKX0DER",
            })
            if r.status_code != 200:
                return []
            data = r.json()
            suggestions = []
            for item in data.get("suggestions", []):
                val = item.get("value")
                if val and val.lower() != query.lower():
                    suggestions.append(val)
            return suggestions[:8]
    except Exception as e:
        print(f"[Amazon Autocomplete] {query}: {e}")
        return []


def classify_amazon_demand(suggestions: list[str]) -> tuple[bool, str]:
    if suggestions:
        return True, f"Amazon suggerisce {len(suggestions)} ricerche correlate — i lettori cercano attivamente in questa nicchia"
    return False, "Nessun suggerimento Amazon per questa query — la nicchia potrebbe essere troppo nuova o servire un angolo diverso"


async def fetch_tiktok_trending(limit: int = 20) -> list[dict]:
    """
    Fetch trending hashtags from TikTok Creative Center's public trend API
    (no auth/cookie required for the anonymous "popular" view). TikTok
    frequently changes or rate-limits this endpoint, so this source degrades
    gracefully to [] on any failure, like YouTube.
    """
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
            "Accept": "application/json, text/plain, */*",
            "Referer": "https://ads.tiktok.com/business/creativecenter/inspiration/popular/hashtag/pc/en",
        }
        async with httpx.AsyncClient(timeout=10, headers=headers) as client:
            r = await client.get("https://ads.tiktok.com/creative_radar_api/v1/popular_trend/hashtag/list", params={
                "page": 1,
                "limit": limit,
                "period": 7,
                "country_code": "US",
                "sort_by": "popular",
            })
            if r.status_code != 200:
                print(f"[TikTok] {r.status_code}: {r.text[:200]}")
                return []
            data = r.json()
            items = []
            for h in data.get("data", {}).get("list", [])[:limit]:
                items.append({
                    "hashtag": h.get("hashtag_name", ""),
                    "rank": h.get("rank"),
                    "publish_cnt": h.get("publish_cnt", 0),
                    "video_views": h.get("video_views", 0),
                })
            return items
    except Exception as e:
        print(f"[TikTok] {e}")
        return []


@app.get("/discover", dependencies=[_AUTH])
async def discover_unbiased(market_language: str = "English"):
    """
    Zero-bias discovery: fetch raw global signals from Reddit + Google Trends,
    then let Claude identify KDP opportunities — no niche pre-selection.
    """
    stamp = now_stamp()

    # Run all fetches in parallel
    reddit_posts, gtrends, youtube_videos, tiktok_hashtags = await asyncio.gather(
        fetch_reddit_global(),
        fetch_google_trending_now(),
        fetch_youtube_trending(),
        fetch_tiktok_trending()
    )

    # Format Reddit data
    reddit_block = ""
    if reddit_posts:
        subs_seen = list(set(p["subreddit"] for p in reddit_posts))
        reddit_block = f"REDDIT — Top & Rising posts TODAY (from {len(subs_seen)} subreddits, no niche filter):\n"
        for i, p in enumerate(reddit_posts[:30], 1):
            snippet = f' | "{p["selftext_snippet"][:80]}..."' if p["selftext_snippet"] else ""
            reddit_block += f'{i}. [r/{p["subreddit"]}] "{p["title"]}" — {p["score"]} upvotes{snippet}\n'
    else:
        reddit_block = "Reddit data unavailable this run.\n"

    # Format Google Trends data
    gtrends_block = ""
    daily = gtrends.get("daily_trending", [])
    realtime = gtrends.get("realtime", [])
    if daily or realtime:
        gtrends_block = "GOOGLE TRENDS — What people are searching RIGHT NOW (no seed, pure signal):\n"
        if daily:
            gtrends_block += f"Today top searches: {', '.join(str(x) for x in daily[:15])}\n"
        if realtime:
            gtrends_block += f"Realtime exploding: {', '.join(str(x) for x in realtime[:10])}\n"
    else:
        gtrends_block = "Google Trends data unavailable this run.\n"

    # Format YouTube Trending data
    youtube_block = ""
    if youtube_videos:
        youtube_block = f"YOUTUBE — Trending NOW ({youtube_videos[0]['region']}, no search query, no niche filter):\n"
        for i, v in enumerate(youtube_videos[:20], 1):
            youtube_block += f'{i}. "{v["title"]}" — {v["channel"]} ({v["views"]:,} views)\n'
    else:
        youtube_block = "YouTube Trending data unavailable (YOUTUBE_API_KEY not set or request failed).\n"

    # Format TikTok Trending data
    tiktok_block = ""
    if tiktok_hashtags:
        tiktok_block = "TIKTOK — Trending hashtags THIS WEEK (no search query, no niche filter):\n"
        for i, h in enumerate(tiktok_hashtags[:20], 1):
            views = f" — {h['video_views']:,} views" if h.get("video_views") else ""
            tiktok_block += f'{i}. #{h["hashtag"]}{views}\n'
    else:
        tiktok_block = "TikTok Trending data unavailable this run.\n"

    lang_cfg = MARKET_LANG_CONFIG.get(market_language, MARKET_LANG_CONFIG["English"])
    amazon_market = lang_cfg["amazon"]
    lang_note = f"\nTARGET MARKET: {market_language} — all book titles and content must be in {market_language}, targeting {amazon_market} readers." if market_language != "English" else ""

    prompt = f"""You are a KDP publishing expert with zero preconceptions about what niche to target.

CURRENT MOMENT: {stamp}{lang_note}

You are about to read RAW, UNFILTERED social data — no topic was specified, no niche was chosen.
Your job: find what is organically emerging and identify KDP book opportunities from it.
{"Translate all book titles and subtitles into " + market_language + " — they must be ready to publish on " + amazon_market + "." if market_language != "English" else ""}

RAW DATA:
{reddit_block}
{gtrends_block}
{youtube_block}
{tiktok_block}

TASK: Identify 5 KDP book opportunities hidden inside this raw data.

RULES — THIS IS CRITICAL:
- You have NO idea what niche to look in — discover it FROM the data
- Each opportunity must come from a REAL signal above (quote it)
- Do NOT apply any pre-existing frameworks about "hot niches"
- Look for: recurring emotional themes, unanswered questions, pain points,
  cultural moments, linguistic patterns across multiple posts
- The niche for each book should emerge ORGANICALLY from patterns you see
- Be specific and surprising — avoid obvious conclusions
- Each of the 5 must be in a DIFFERENT niche/category

TITLE RULES — Amazon KDP rejects listings with "title": "subtitle"-style titles
or keyword-stuffed titles (combined title+subtitle over ~200 characters):
- "title": SHORT and punchy, under 60 characters, NO colon followed by a long
  second subtitle baked into it
- "subtitle": separate field for SEO/keywords, under 140 characters
- title + subtitle combined MUST be under 200 characters total

KDP FIT — for each opportunity, honestly assess:
- "kdp_fit": "high" | "medium" | "low" — how well this niche translates into a
  book/journal/workbook/planner people would actually PAY for, vs. a topic
  people only consume as free video/news/forum content
  - "high" = clear evergreen audience for a companion book/journal/workbook/planner
  - "medium" = possible but needs the right angle (e.g. fandom journal, not a "guide")
  - "low" = pure entertainment/news/gaming content — audience expects free video/wiki, not a paid book
- "kdp_fit_reason": 1 sentence explaining the verdict, specific to this niche

CROSS-SOURCE SCORING — for each opportunity:
- "sources": list which raw data blocks above (reddit, google, youtube, tiktok) contain
  a signal supporting this opportunity. Only include a source if you can quote
  a real signal from it in "data_signals".
- "stage": classify based on signal strength and cross-source presence:
  - "Esplode" = appears in 3-4 sources OR extremely strong signal in 2 (breaking out NOW)
  - "Forte" = appears in 2 sources with strong signal (high demand, validate competition)
  - "Crescita" = appears in 1-2 sources, growing but not yet saturated (optimal KDP window)
  - "Pre-virale" = weak/early signal in 1 source — gap opportunity but unproven demand

Return ONLY raw JSON, no markdown, ASCII-safe strings only:
{{"opportunities":[
{{
  "niche": "The organic niche you discovered (do not use predefined categories)",
  "pattern": "The specific pattern you noticed across multiple data points",
  "data_signals": ["exact Reddit post title or Google query or YouTube title 1", "signal 2", "signal 3"],
  "sources": ["reddit","google"],
  "stage": "Crescita",
  "kdp_fit": "high",
  "kdp_fit_reason": "1 sentence on why this fits a paid book/journal/workbook audience",
  "heat": 4,
  "why_now": "1 sentence: why this moment is the right time for this book",
  "books": [
    {{"type":"Book type","title":"Specific title","subtitle":"Amazon SEO subtitle"}},
    {{"type":"Book type 2","title":"Title 2","subtitle":"Subtitle 2"}}
  ]
}},
{{
  "niche": "Discovered niche 2",
  "pattern": "Pattern observed",
  "data_signals": ["signal 1","signal 2"],
  "sources": ["reddit"],
  "stage": "Pre-virale",
  "kdp_fit": "medium",
  "kdp_fit_reason": "1 sentence",
  "heat": 5,
  "why_now": "1 sentence",
  "books": [
    {{"type":"TYPE","title":"TITLE","subtitle":"subtitle"}},
    {{"type":"TYPE","title":"TITLE 2","subtitle":"subtitle 2"}}
  ]
}},
{{
  "niche": "Discovered niche 3",
  "pattern": "Pattern",
  "data_signals": ["signal 1","signal 2"],
  "sources": ["google","youtube"],
  "stage": "Forte",
  "kdp_fit": "low",
  "kdp_fit_reason": "1 sentence",
  "heat": 3,
  "why_now": "1 sentence",
  "books": [
    {{"type":"TYPE","title":"TITLE","subtitle":"subtitle"}},
    {{"type":"TYPE","title":"TITLE 2","subtitle":"subtitle 2"}}
  ]
}},
{{
  "niche": "Discovered niche 4",
  "pattern": "Pattern",
  "data_signals": ["signal 1","signal 2"],
  "sources": ["reddit","google","youtube"],
  "stage": "Esplode",
  "kdp_fit": "high",
  "kdp_fit_reason": "1 sentence",
  "heat": 4,
  "why_now": "1 sentence",
  "books": [
    {{"type":"TYPE","title":"TITLE","subtitle":"subtitle"}},
    {{"type":"TYPE","title":"TITLE 2","subtitle":"subtitle 2"}}
  ]
}},
{{
  "niche": "Discovered niche 5",
  "pattern": "Pattern",
  "data_signals": ["signal 1","signal 2"],
  "sources": ["reddit"],
  "stage": "Crescita",
  "kdp_fit": "medium",
  "kdp_fit_reason": "1 sentence",
  "heat": 5,
  "why_now": "1 sentence",
  "books": [
    {{"type":"TYPE","title":"TITLE","subtitle":"subtitle"}},
    {{"type":"TYPE","title":"TITLE 2","subtitle":"subtitle 2"}}
  ]
}}
]}}"""

    try:
        text = await call_claude(prompt, 3500)
        result = parse_json_safe(text)

        # Gap analysis — cross-check each discovered niche against Amazon book count
        # and Amazon search autocomplete (real reader demand signal).
        # Run sequentially with a small delay: 10 concurrent requests to Amazon
        # from the same datacenter IP gets rate-limited/blocked almost every time.
        import urllib.parse
        opportunities = result.get("opportunities", [])
        for idx, o in enumerate(opportunities):
            niche = o.get("niche", "")
            if idx > 0:
                await asyncio.sleep(0.8)
            gr, suggestions = await asyncio.gather(
                fetch_amazon_book_count(niche),
                fetch_amazon_autocomplete(niche),
            )
            search_url = "https://www.amazon.com/s?" + urllib.parse.urlencode({"k": niche, "i": "stripbooks"})
            count = gr.get("count")
            level, note = classify_amazon_saturation(count)
            o["gap_analysis"] = {
                "amazon_results": count,
                "saturation": level,
                "note": note,
                "query": gr.get("query"),
                "search_url": search_url,
            }
            has_demand, demand_note = classify_amazon_demand(suggestions)
            o["amazon_demand"] = {
                "suggestions": suggestions,
                "has_signal": has_demand,
                "note": demand_note,
                "search_url": search_url,
            }

        result["meta"] = {
            "mode": "zero_bias_discovery",
            "reddit_posts_analyzed": len(reddit_posts),
            "reddit_subreddits": list(set(p["subreddit"] for p in reddit_posts))[:10],
            "google_daily_trending": gtrends.get("daily_trending", [])[:10],
            "google_realtime": gtrends.get("realtime", [])[:5],
            "youtube_videos_analyzed": len(youtube_videos),
            "youtube_configured": bool(YOUTUBE_API_KEY),
            "tiktok_hashtags_analyzed": len(tiktok_hashtags),
            "fetched_at": stamp,
            "note": "No niche was specified — opportunities discovered purely from raw data"
        }
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/positioning", dependencies=[_AUTH])
async def positioning_to_books(req: PositioningRequest):
    """
    Inverse flow from /discover: instead of starting from 'what's trending'
    and writing a book toward it, start from the author's deliberate
    editorial positioning (identity, ideal reader psychographics, promised
    transformation, competitive differentiation, unfair advantage) and derive
    the niche, angle, and book concepts that express it.
    """
    stamp = now_stamp()

    prompt = f"""You are a KDP positioning strategist helping a publisher who has
ALREADY done market analysis and decided their editorial positioning — they are
NOT chasing trends. Your job is to translate their positioning into a clear
niche/angle and concrete book concepts that express it.

CURRENT MOMENT: {stamp}

AUTHOR/PUBLISHER POSITIONING BRIEF:
- Identity & desired perception: {req.author_identity}
- Ideal reader (psychographic, not demographic): {req.ideal_reader}
- Promised transformation: {req.transformation}
- Competitors & differentiation: {req.competitors or "not specified"}
- Unfair advantage (unique knowledge/experience): {req.unfair_advantage or "not specified"}

TASK:
1. Derive ONE clear editorial positioning statement that ties identity, reader,
   transformation and differentiation together into a coherent angle.
2. Derive the niche this positioning lives in (specific, not generic).
3. Propose 3 book concepts that EXPRESS this positioning — each a different
   format/angle on the same positioning, not random unrelated ideas.

TITLE RULES — Amazon KDP rejects listings with "title": "subtitle"-style titles
or keyword-stuffed titles (combined title+subtitle over ~200 characters):
- "title": SHORT and punchy, under 60 characters, NO colon followed by a long
  second subtitle baked into it
- "subtitle": separate field for SEO/keywords, under 140 characters
- title + subtitle combined MUST be under 200 characters total

Return ONLY raw JSON, no markdown, ASCII-safe strings only:
{{"positioning_summary":"2-3 sentences tying identity, reader, transformation and differentiation into one coherent editorial angle",
"niche":"the specific niche this positioning lives in",
"why_this_works":"1-2 sentences on why this positioning is differentiated and defensible vs competitors",
"books":[
  {{"type":"Book type","title":"Short punchy title","subtitle":"SEO subtitle expressing the positioning"}},
  {{"type":"Book type 2","title":"Title 2","subtitle":"Subtitle 2"}},
  {{"type":"Book type 3","title":"Title 3","subtitle":"Subtitle 3"}}
]}}"""

    try:
        text = await call_claude(prompt, 2000)
        result = parse_json_safe(text)
        result["meta"] = {
            "mode": "positioning",
            "fetched_at": stamp,
        }
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/trends", dependencies=[_AUTH])
async def get_trends(req: TrendRequest):
    year = datetime.now().year
    stamp = now_stamp()
    platforms_str = ", ".join(req.platforms)

    lang = req.market_language or "English"
    lang_cfg = MARKET_LANG_CONFIG.get(lang, MARKET_LANG_CONFIG["English"])
    amazon_market = lang_cfg["amazon"]
    lang_subreddits = lang_cfg["subreddits"]

    lang_code = LANG_CODE_MAP.get(lang, "en")
    tf_info = get_timeframe(req.timeframe or "week")

    # Primary: Google + YouTube autocomplete (language-aware, works from any IP)
    # Secondary: Google Trends
    # Reddit: English only (low signal for non-EN markets)
    fetch_tasks = [
        fetch_multi_autocomplete(req.niche, lang_code),
        fetch_google_trends(req.niche, req.keyword or "", req.timeframe or "week"),
    ]
    if lang == "English":
        fetch_tasks.append(fetch_reddit_posts(req.niche, req.keyword or "", req.timeframe or "week", force_subreddits=lang_subreddits))
    results = await asyncio.gather(*fetch_tasks, return_exceptions=True)

    autocomplete_data = results[0] if isinstance(results[0], dict) else {"google": [], "youtube": []}
    gtrends = results[1] if isinstance(results[1], dict) else {}
    reddit_posts = results[2] if len(results) > 2 and isinstance(results[2], list) else []

    autocomplete_block = ""
    if autocomplete_data["google"]:
        autocomplete_block += f"GOOGLE SEARCH — What {lang} speakers actively search in this niche:\n"
        for s in autocomplete_data["google"][:20]:
            autocomplete_block += f'  - "{s}"\n'
    if autocomplete_data["youtube"]:
        autocomplete_block += f"\nYOUTUBE — What {lang} speakers want to learn (video intent):\n"
        for s in autocomplete_data["youtube"][:12]:
            autocomplete_block += f'  - "{s}"\n'
    if not autocomplete_block:
        autocomplete_block = "Autocomplete data unavailable this run.\n"

    gtrends_summary = ""
    if gtrends.get("avg_interest"):
        gtrends_summary = "GOOGLE TRENDS — Interest over time:\n"
        for term, val in gtrends["avg_interest"].items():
            gtrends_summary += f"  '{term}': {val}/100\n"
        if gtrends.get("rising_queries"):
            gtrends_summary += "Rising queries:\n"
            for term, queries in gtrends["rising_queries"].items():
                if queries:
                    tops = [q.get("query","") for q in queries[:3]]
                    gtrends_summary += f"  Under '{term}': {', '.join(tops)}\n"

    reddit_summary = ""
    if reddit_posts:
        reddit_summary = "\nREDDIT (English signal) — Viral posts:\n"
        for i, p in enumerate(reddit_posts[:12], 1):
            reddit_summary += f'{i}. [r/{p["subreddit"]}] "{p["title"]}" — {p["score"]} upvotes\n'

    prompt = f"""You are a KDP publishing expert analyzing REAL search intent data.

CURRENT MOMENT: {stamp}
ANALYSIS WINDOW: {tf_info["label"]}
STRATEGIC CONTEXT: {tf_info["strategy"]}
TARGET MARKET: {lang} — books must be written in {lang}, targeting {amazon_market}
{"IMPORTANT: All book titles, subtitles, descriptions must be in " + lang + ". Trends must resonate with " + lang + "-speaking readers and their cultural context." if lang != "English" else ""}

REAL SEARCH INTENT DATA (what people are actively searching right now):
{autocomplete_block}
{gtrends_summary}{reddit_summary}

TASK: Identify 4 SPECIFIC, UNDERSERVED KDP book opportunities for the "{req.niche}" category on {amazon_market}.
Platforms context: {platforms_str}
{f'Keyword focus: "{req.keyword}"' if req.keyword else ''}

INTERPRETATION GUIDE:
- Autocomplete queries = proven demand (people typing this = they want content on it)
- High Google Trends score = broad awareness
- Rising queries = emerging, not yet saturated — BEST opportunity
- Find the intersection: high search intent + low existing books = gap

UNIQUENESS RULES:
- Do NOT suggest: generic journals, gratitude journals, mindfulness basics, morning routines
- Each trend must cite a SPECIFIC autocomplete query from the data above as its signal
- Find angles that feel fresh and slightly unexpected

TITLE RULES:
- "title": SHORT and punchy, under 60 characters, NO colon-subtitle inside it
- "subtitle": separate SEO field, under 140 characters
- title + subtitle combined MUST be under 200 characters

Return ONLY raw JSON, no markdown, ASCII-safe strings only:
{{"trends":[
{{"name":"SPECIFIC TREND","platform":"Google/YouTube","description":"2 sentences citing a specific autocomplete query","heat":4,"data_signal":"Exact autocomplete query that proves demand","books":[
{{"type":"TYPE","title":"TITLE","subtitle":"SEO subtitle"}},
{{"type":"TYPE","title":"TITLE 2","subtitle":"SEO subtitle 2"}}
]}},
{{"name":"TREND 2","platform":"Google/YouTube","description":"2 sentences","heat":5,"data_signal":"exact signal","books":[
{{"type":"TYPE","title":"TITLE","subtitle":"subtitle"}},
{{"type":"TYPE","title":"TITLE 2","subtitle":"subtitle 2"}}
]}},
{{"name":"TREND 3","platform":"Google/YouTube","description":"2 sentences","heat":3,"data_signal":"signal","books":[
{{"type":"TYPE","title":"TITLE","subtitle":"subtitle"}},
{{"type":"TYPE","title":"TITLE 2","subtitle":"subtitle 2"}}
]}},
{{"name":"TREND 4","platform":"Google/YouTube","description":"2 sentences","heat":4,"data_signal":"signal","books":[
{{"type":"TYPE","title":"TITLE","subtitle":"subtitle"}},
{{"type":"TYPE","title":"TITLE 2","subtitle":"subtitle 2"}}
]}}
]}}"""

    try:
        text = await call_claude(prompt, 3000)
        result = parse_json_safe(text)
        result["meta"] = {
            "reddit_posts_found": len(reddit_posts),
            "google_autocomplete_found": len(autocomplete_data["google"]),
            "youtube_autocomplete_found": len(autocomplete_data["youtube"]),
            "gtrends_avg": gtrends.get("avg_interest", {}),
            "fetched_at": stamp,
            "timeframe": req.timeframe or "week",
            "timeframe_label": tf_info["label"],
            "timeframe_strategy": tf_info["strategy"],
            "data_sources": {
                "google_autocomplete": len(autocomplete_data["google"]) > 0,
                "youtube_autocomplete": len(autocomplete_data["youtube"]) > 0,
                "reddit": len(reddit_posts) > 0,
                "google_trends": bool(gtrends.get("avg_interest"))
            },
            # Pass niche-specific search queries to the frontend for prominent display
            "niche_queries": {
                "google": autocomplete_data.get("google", [])[:20],
                "youtube": autocomplete_data.get("youtube", [])[:12],
            }
        }
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/validate-niche", dependencies=[_AUTH])
async def validate_niche(req: dict):
    """Quick validation: fetch autocomplete + Google Trends for a specific niche name."""
    niche = req.get("niche", "")
    market_language = req.get("market_language", "English")
    lang_code = LANG_CODE_MAP.get(market_language, "en")
    lang_cfg = MARKET_LANG_CONFIG.get(market_language, MARKET_LANG_CONFIG["English"])
    amazon_market = lang_cfg["amazon"]
    import urllib.parse

    autocomplete_data = await fetch_multi_autocomplete(niche, lang_code)

    search_url = "https://" + lang_cfg["amazon"].split("/")[0].replace("Amazon.", "amazon.") + "/s?" + urllib.parse.urlencode({"k": niche, "i": "stripbooks"})
    gtrends_url = "https://trends.google.com/trends/explore?" + urllib.parse.urlencode({"q": niche, "geo": lang_code.upper()})

    return {
        "niche": niche,
        "market_language": market_language,
        "amazon_market": amazon_market,
        "google_suggestions": autocomplete_data.get("google", [])[:15],
        "youtube_suggestions": autocomplete_data.get("youtube", [])[:10],
        "google_trends_score": None,
        "trend_direction": None,
        "rising_queries": [],
        "amazon_search_url": search_url,
        "google_trends_url": gtrends_url,
    }


@app.post("/trending-now", dependencies=[_AUTH])
async def trending_now(req: dict):
    """Return what's viral right now on Google in the target market.
    Powered by Apify automation-lab/google-trends-scraper trending topics."""
    market_language = req.get("market_language", "English")
    data = await fetch_google_trends_apify("", "", market_language)
    pool = []
    for qs in (data.get("rising_queries") or {}).values():
        for q in qs:
            pool.append({"keyword": q.get("query", ""), "traffic": q.get("value", "")})
    return {
        "market_language": market_language,
        "geo": data.get("geo", "US"),
        "trending": pool[:20],
    }


@app.post("/niches", dependencies=[_AUTH])
async def discover_niches(req: NicheRequest):
    stamp = now_stamp()
    year = datetime.now().year

    tf_info = get_timeframe(req.timeframe or "week")
    general_posts, gtrends = await asyncio.gather(
        fetch_reddit_posts("Any niche", req.keyword or "", req.timeframe or "week"),
        fetch_google_trends("Any niche", req.keyword or "", req.timeframe or "week")
    )

    reddit_summary = ""
    if general_posts:
        subs_used = list(set(p["subreddit"] for p in general_posts))
        reddit_summary = f"Reddit posts this week (subreddits: {', '.join(subs_used[:6])}):\n"
        for p in general_posts[:20]:
            reddit_summary += f"- [r/{p['subreddit']}] \"{p['title']}\" ({p['score']} pts)\n"

    gtrends_summary = ""
    if gtrends.get("rising_queries"):
        gtrends_summary = f"Google rising queries (seeds: {', '.join(gtrends.get('terms',[]))}):\n"
        for term, queries in gtrends["rising_queries"].items():
            for q in queries[:3]:
                gtrends_summary += f"  - {q.get('query','')}\n"

    prompt = f"""You are a KDP niche analyst. Today is {stamp}.

REAL DATA:
{reddit_summary}
{gtrends_summary}

Find 8 SPECIFIC, UNDERSERVED KDP niches that are genuinely emerging RIGHT NOW.

RULES:
- NO generic suggestions (no "self-help journal", "gratitude journal", "mindfulness basics")
- Each niche must be supported by a specific signal from the data
- Look for niches that are JUST emerging, not already saturated
- Think cross-cultural, counter-intuitive, niche-within-niche angles
- Variety: each of the 8 must be in a different micro-category

Return ONLY raw JSON, ASCII-safe text only:
{{"niches":[
{{"name":"Specific 2-5 word niche","reason":"Why emerging now — 1 sentence referencing data","heat":4,"example_book":"Specific KDP title","data_signal":"Exact Reddit post or Google query that supports this"}},
{{"name":"Niche 2","reason":"1 sentence","heat":5,"example_book":"Title","data_signal":"signal"}},
{{"name":"Niche 3","reason":"1 sentence","heat":3,"example_book":"Title","data_signal":"signal"}},
{{"name":"Niche 4","reason":"1 sentence","heat":5,"example_book":"Title","data_signal":"signal"}},
{{"name":"Niche 5","reason":"1 sentence","heat":4,"example_book":"Title","data_signal":"signal"}},
{{"name":"Niche 6","reason":"1 sentence","heat":3,"example_book":"Title","data_signal":"signal"}},
{{"name":"Niche 7","reason":"1 sentence","heat":4,"example_book":"Title","data_signal":"signal"}},
{{"name":"Niche 8","reason":"1 sentence","heat":5,"example_book":"Title","data_signal":"signal"}}
]}}"""

    try:
        text = await call_claude(prompt, 2500)
        result = parse_json_safe(text)
        result["meta"] = {
            "reddit_posts_found": len(general_posts),
            "gtrends_seeds": gtrends.get("terms",[]),
            "fetched_at": stamp
        }
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/generate", dependencies=[_AUTH])
async def generate_content(req: GenerateRequest):
    book_ctx = (
        f'Book: "{req.book_title}" - {req.book_subtitle}\n'
        f'Type: {req.book_type}\n'
        f'Trend: "{req.trend_name}"\n'
        f'Audience: {req.audience}'
    )
    # Get book-type-specific template
    book_tmpl = get_book_template(req.book_type)
    book_format_ctx = f"""BOOK FORMAT RULES (follow precisely for {req.book_type}):
Structure: {book_tmpl["structure"]}
Chapter instruction: {book_tmpl["chapter_instruction"]}
Length note: {book_tmpl["length_note"]}
Unique elements: {book_tmpl["unique_elements"]}"""

    voice_ctx = build_voice_ctx(
        tone=req.tone,
        language=req.language or "English",
        cultural_inspiration=req.cultural_inspiration or "",
        chapter_length=req.chapter_length or "medium",
        reader_persona=req.reader_persona or "",
        custom_instructions=req.custom_instructions or ""
    )
    kw_ctx = ""
    if req.amazon_keywords:
        kw_ctx = (
            f"\nKDP KEYWORD OPTIMIZATION — These are real Amazon search terms buyers use. "
            f"Weave them naturally into chapter titles, section headings, and the first/last paragraph of each chapter "
            f"(never stuff them unnaturally): {', '.join(req.amazon_keywords[:12])}\n"
        )

    niche_ctx = ""
    if req.niche_analysis:
        na = req.niche_analysis
        parts = []
        if na.get("best_angle"):
            parts.append(f"UNEXPLOITED ANGLE (use this as the core POV of the book): {na['best_angle']}")
        if na.get("opportunities"):
            opps = na["opportunities"] if isinstance(na["opportunities"], list) else []
            if opps:
                parts.append("OPPORTUNITIES TO EXPLOIT:\n" + "\n".join(f"  - {o}" for o in opps[:4]))
        if na.get("keyword_gems"):
            gems = na["keyword_gems"] if isinstance(na["keyword_gems"], list) else []
            if gems:
                parts.append(f"HIGH-VALUE KEYWORDS (integrate naturally): {', '.join(str(g) for g in gems[:6])}")
        if na.get("risks"):
            risks = na["risks"] if isinstance(na["risks"], list) else []
            if risks:
                parts.append("RISKS TO AVOID/ADDRESS: " + "; ".join(str(r) for r in risks[:2]))
        if parts:
            niche_ctx = "\nNICHE INTELLIGENCE (apply this to sharpen focus and differentiation):\n" + "\n".join(parts) + "\n"
    length_map = {"short": "800-1000", "medium": "1200-1600", "long": "2000-2500"}
    word_count = length_map.get(req.chapter_length or "medium", "1200-1600")

    if req.tab == "outline":
        prompt = f"""You are a bestselling KDP author. Create a detailed book outline.

{book_ctx}
{kw_ctx}{niche_ctx}
{book_format_ctx}

{voice_ctx}

Write a professional outline with 10 chapters structured for a {req.book_type}.
For each chapter: chapter number, punchy title, and 4 subsection titles.

Format:
Chapter 1: Title
  1.1 Subsection
  1.2 Subsection
  1.3 Subsection
  1.4 Subsection

All 10 chapters. No extra text."""
        max_tok = 3000

    elif req.tab in ("chapter", "allchapters"):
        n = req.chapter_num or 1
        prompt = f"""You are a bestselling KDP author. Write Chapter {n}.

{book_ctx}
{kw_ctx}{niche_ctx}
{f"Outline:{chr(10)}{req.outline[:600]}" if req.outline else ""}

{book_format_ctx}

{voice_ctx}

Write CHAPTER {n} completely following the book format rules above.
- Chapter title as header
- Follow the opening style specified above EXACTLY
- Follow the chapter instruction for {req.book_type} precisely
- Target: {word_count} words

Make it feel like no other chapter in any other book."""
        max_tok = 8000

    elif req.tab == "intro":
        prompt = f"""You are a bestselling KDP author. Write the Introduction AND Conclusion.

{book_ctx}
{kw_ctx}
{voice_ctx}

INTRODUCTION ({word_count} words):
- Follow the opening style specified above
- Why this book exists and who it is for
- What the reader will gain
- How to use the book

CONCLUSION (600-800 words):
- Recap the transformation arc
- Motivational, specific closing
- Next steps / call to action"""
        max_tok = 6000

    elif req.tab == "full":
        prompt = f"""You are a bestselling KDP author. Write a complete {req.book_type}.

{book_ctx}
{kw_ctx}
{voice_ctx}

Write the FULL book:
- Introduction (400-500 words)
- 6 complete chapters ({word_count} words each) with content, prompts, exercises
- Conclusion (300-400 words)

Follow voice guidelines throughout. Make every chapter feel distinct."""
        max_tok = 8000
    else:
        raise HTTPException(status_code=400, detail=f"Unknown tab: {req.tab}")

    try:
        allow_trunc = req.tab in ("chapter", "allchapters", "draft", "intro", "full")
        text = await call_claude(prompt, max_tok, allow_truncated=allow_trunc)
        return {"content": text, "tab": req.tab, "chapter_num": req.chapter_num}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/generate-all", dependencies=[_AUTH])
async def generate_all_chapters(req: AllChaptersRequest):
    import re as _re
    book_ctx = (
        f'Book: "{req.book_title}" - {req.book_subtitle}\n'
        f'Type: {req.book_type}\n'
        f'Trend: "{req.trend_name}"\n'
        f'Audience: {req.audience}'
    )
    length_map = {"short": "800-1000", "medium": "1200-1600", "long": "2000-2500"}
    word_count = length_map.get(req.chapter_length or "medium", "1200-1600")

    chapter_pattern = _re.compile(r'^(?:Chapter\s+)?(\d+)[:\.\)]\s*(.+)$', _re.MULTILINE | _re.IGNORECASE)
    chapters = []
    for m in chapter_pattern.finditer(req.outline):
        if '.' not in m.group(1):
            chapters.append({"num": int(m.group(1)), "title": m.group(2).strip()})
    seen_nums, unique_chapters = set(), []
    for ch in sorted(chapters, key=lambda x: x["num"]):
        if ch["num"] not in seen_nums:
            seen_nums.add(ch["num"])
            unique_chapters.append(ch)
    if not unique_chapters:
        raise HTTPException(status_code=400, detail="Nessun capitolo trovato. Genera prima l'Outline.")
    _MAX_CHAPTERS = 25
    if len(unique_chapters) > _MAX_CHAPTERS:
        raise HTTPException(status_code=400, detail=f"Troppi capitoli: max {_MAX_CHAPTERS} per richiesta.")

    total = len(unique_chapters)
    outline_snippet = req.outline[:800]

    kw_ctx = ""
    if req.amazon_keywords:
        kw_ctx = (
            f"\nKDP KEYWORD OPTIMIZATION — These are real Amazon search terms buyers use. "
            f"Weave them naturally into chapter titles, section headings, and the first/last paragraph of each chapter "
            f"(never stuff them unnaturally): {', '.join(req.amazon_keywords[:12])}\n"
        )

    start_from = max(1, req.start_from or 1)

    async def chapter_stream():
        for ch in unique_chapters:
            n, title = ch["num"], ch["title"]
            # Skip already-generated chapters when resuming
            if n < start_from:
                yield json.dumps({
                    "chapter": n, "title": title, "content": "",
                    "skipped": True, "total": total, "done": False
                }, ensure_ascii=False) + "\n"
                continue
            # Each chapter gets its own random opening style + uniqueness seed
            voice_ctx = build_voice_ctx(
                tone=req.tone,
                language=req.language or "English",
                cultural_inspiration=req.cultural_inspiration or "",
                chapter_length=req.chapter_length or "medium",
                reader_persona=req.reader_persona or "",
                custom_instructions=req.custom_instructions or ""
            )
            na_ctx = ""
            if req.niche_analysis:
                na = req.niche_analysis
                angle = na.get("best_angle", "")
                opps = na.get("opportunities", [])
                gems = na.get("keyword_gems", [])
                na_parts = []
                if angle: na_parts.append(f"CORE ANGLE: {angle}")
                if opps:  na_parts.append("OPPORTUNITIES: " + "; ".join(str(o) for o in opps[:3]))
                if gems:  na_parts.append(f"POWER KEYWORDS: {', '.join(str(g) for g in gems[:5])}")
                if na_parts:
                    na_ctx = "\nNICHE INTELLIGENCE:\n" + "\n".join(na_parts) + "\n"

            prompt = f"""You are a bestselling KDP author. Write Chapter {n} of {total}.

{book_ctx}
{kw_ctx}{na_ctx}
Chapter title: {title}
Outline reference:
{outline_snippet}

{voice_ctx}

Write CHAPTER {n} — "{title}" — completely:
- Open with chapter title as header
- Follow the opening style above EXACTLY — make it vivid and unexpected
- 3-4 full sections with subheadings
- Content, prompts or exercises appropriate for {req.book_type}
- Chapter summary at the end
- Target: {word_count} words

This chapter must feel completely distinct from chapters before it.
Use the uniqueness seed to ensure this version is unlike any previous generation."""

            try:
                text = await call_claude(prompt, 8000, allow_truncated=True)
                yield json.dumps({
                    "chapter": n, "title": title, "content": text,
                    "total": total, "done": False
                }, ensure_ascii=False) + "\n"
            except Exception as e:
                yield json.dumps({
                    "chapter": n, "title": title, "content": "",
                    "error": str(e), "total": total, "done": False
                }) + "\n"

            await asyncio.sleep(0.5)

        yield json.dumps({"done": True, "total": total}) + "\n"

    return StreamingResponse(
        chapter_stream(),
        media_type="application/x-ndjson",
        headers={"X-Total-Chapters": str(total)}
    )


@app.post("/analyze-market", dependencies=[_AUTH])
async def analyze_market(req: dict):
    """
    Analyze market positioning based on user-pasted Amazon competitor data.
    User pastes titles/prices from Amazon search results, Claude analyzes.
    """
    titles = req.get("titles", [])         # list of competitor titles
    prices = req.get("prices", [])          # list of prices as strings
    book_title = req.get("book_title", "")
    book_type = req.get("book_type", "")
    niche = req.get("niche", "")
    stamp = now_stamp()

    if not titles:
        raise HTTPException(status_code=400, detail="Provide at least 3 competitor titles")

    competitors_block = "\n".join(
        f'- "{t}" — ${prices[i]}' if i < len(prices) and prices[i] else f'- "{t}"'
        for i, t in enumerate(titles)
    )

    prompt = f"""You are an Amazon KDP market analyst. Analyze this competitor data and give strategic advice.
IMPORTANT: Detect the language of the titles/niche provided and write ALL text fields in that same language.

My book: "{book_title}" ({book_type}) in niche: "{niche}"
Analysis date: {stamp}

COMPETITOR TITLES ON AMAZON:
{competitors_block}

Analyze and return ONLY raw JSON, ASCII-safe:
{{"analysis":{{
  "competitor_count": {len(titles)},
  "saturation_level": "low|medium|high",
  "saturation_score": 3,
  "avg_price_estimate": "$X.XX",
  "price_sweet_spot": "$X.XX - $X.XX",
  "title_patterns": ["pattern 1 seen in competitors","pattern 2","pattern 3"],
  "gaps_identified": ["gap or angle not covered by competitors 1","gap 2","gap 3"],
  "positioning_advice": "2-3 sentences on how to differentiate from these competitors",
  "recommended_title_angle": "Specific angle for MY book title that stands out from these competitors",
  "subtitle_keywords": ["keyword 1","keyword 2","keyword 3","keyword 4"],
  "verdict": "go|caution|avoid",
  "verdict_reason": "1 sentence explaining the verdict"
}}}}"""

    try:
        text = await call_claude(prompt, 1500)
        result = parse_json_safe(text)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/title-variants", dependencies=[_AUTH])
async def generate_title_variants(req: dict):
    """
    DLAB Title Generator — 7 positioning strategies × 8 proven schemas.
    Each title scored on 5 evaluation criteria.
    """
    book_type = req.get("book_type", "")
    niche = req.get("niche", "")
    trend = req.get("trend", "")
    audience = req.get("audience", "")
    current_title = req.get("current_title", "")
    language = req.get("language", "English")
    real_keywords = req.get("real_keywords", [])
    usp = req.get("usp", "")
    avatar_summary = req.get("avatar_summary", "")
    preferred_strategies = req.get("preferred_strategies", [])

    kw_ctx = f"\nReal Amazon/Google search terms (weave into titles naturally): {', '.join(real_keywords[:15])}" if real_keywords else ""
    usp_ctx = f"\nUSP / unique feature: {usp}" if usp else ""
    avatar_ctx = f"\nAvatar summary: {avatar_summary}" if avatar_summary else ""
    strategy_ctx = f"\nPreferred strategies (prioritize these): {', '.join(preferred_strategies)}" if preferred_strategies else ""

    prompt = f"""You are a KDP title strategist using the DLAB framework. Generate exactly 5 title + subtitle combinations.

Book type: {book_type}
Niche/trend: {niche} — {trend}
Audience: {audience}
Working title: "{current_title}"{usp_ctx}{avatar_ctx}{kw_ctx}{strategy_ctx}
Output language: {language}

POSITIONING STRATEGIES available (pick the 5 most effective for this specific niche/audience):
- Reframe: redefine a common belief about the topic
- Contrarian: take the opposite position from all competitors
- USP: lead with a specific concrete feature others don't have
- Blue Ocean: target an underserved sub-segment to avoid direct competition
- Authority: imply exclusive insider knowledge or a clinical/professional method
- Story-Driven: frame as a transformation journey (from X to Y)
- Process-Driven: a clear named method or system (3 steps, formula, blueprint)

PROVEN SCHEMAS to apply (use different ones for variety):
1. [Keyword] per [Specific Audience]
2. [Number] [methods/secrets/strategies] per [achieve goal]
3. Come [achieve goal] senza [common obstacle]
4. [Benefit], [Keyword]
5. [Achieve goal] in [time period]
6. [Action verb] [keyword] per [benefit]
7. [Keyword]: La Guida Completa / Il Blueprint / I Segreti
8. Contrarian or Authority angle title

EVALUATION — score each title 1-10 on:
- chiarezza: does a browser instantly understand what it's about?
- memorabilita: easy to say, remember, and search?
- curiosita: creates desire to know more?
- rilevanza: matches what the target reader is actively searching for?
- competizione: stands out from existing Amazon titles in this niche?

Return ONLY valid JSON:
{{"variants":[
  {{
    "strategy": "Strategy name",
    "schema": "Schema used (brief description)",
    "title": "Main title",
    "subtitle": "Subtitle (include keywords, USP, audience where natural)",
    "why": "One sentence: why this strategy+schema works for this specific audience",
    "scores": {{"chiarezza":8,"memorabilita":7,"curiosita":9,"rilevanza":8,"competizione":7}},
    "total": 39
  }}
]}}

Generate exactly 5 variants using 5 DIFFERENT strategies. No duplicate strategies."""

    try:
        text = await call_claude(prompt, 2500)
        return parse_json_safe(text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/cover-brief", dependencies=[_AUTH])
async def generate_cover_brief(req: dict):
    """Generate a cover design brief for Canva/Midjourney based on title + positioning."""
    title = req.get("title", "")
    subtitle = req.get("subtitle", "")
    niche = req.get("niche", "")
    book_type = req.get("book_type", "")
    audience = req.get("audience", "")
    usp = req.get("usp", "")
    strategy = req.get("strategy", "")
    language = req.get("language", "English")

    prompt = f"""You are an expert book cover designer and KDP publishing strategist.
Generate a complete cover design brief for this book.

Title: "{title}"
Subtitle: "{subtitle}"
Niche: {niche}
Book type: {book_type}
Target audience: {audience}
USP: {usp}
Positioning strategy: {strategy}
Language/market: {language}

IMPORTANT: Detect the language of the title/niche and write ALL text fields in that same language.

A great cover must:
1. Communicate the emotional promise instantly (primordial emotion)
2. Target the avatar visually — they should see themselves in it
3. Be clear and decisive — no confusion about topic
4. Use colors that match the emotional tone of the content
5. Have a title that is readable as a thumbnail (small on mobile)

Return ONLY valid JSON:
{{
  "mood": "overall visual mood in 3-5 words",
  "primary_emotion": "the one emotion the cover must evoke",
  "color_palette": {{
    "primary": "#hexcode — name and why",
    "secondary": "#hexcode — name and why",
    "accent": "#hexcode — name and why",
    "background": "#hexcode — name and why"
  }},
  "typography": {{
    "title_style": "font style description (e.g. bold serif, minimal sans-serif, handwritten)",
    "title_size": "dominant/medium/subtle",
    "subtitle_style": "description",
    "recommendation": "specific font pairing recommendation"
  }},
  "visual_elements": ["element 1 with placement", "element 2", "element 3"],
  "composition": "brief description of layout and hierarchy",
  "avoid": ["thing to avoid 1", "thing to avoid 2", "thing to avoid 3"],
  "canva_search": ["search term 1 for Canva templates", "search term 2", "search term 3"],
  "midjourney_prompt": "complete ready-to-use Midjourney prompt for the cover image (no text, just background/visual)",
  "thumbnail_test": "how this cover will look at 80x120px — what survives and what gets lost"
}}"""

    try:
        text = await call_claude(prompt, 2000)
        return parse_json_safe(text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# Real Amazon KDP browse categories + BISAC codes by book type
_KDP_CATEGORIES = {
    "Guided Journal": {
        "categories": ["Books > Self-Help > Journaling & Bullet Journaling", "Books > Self-Help > Happiness"],
        "bisac": ["SEL031000 SELF-HELP / Journaling", "SEL016000 SELF-HELP / Happiness"],
    },
    "Workbook": {
        "categories": ["Books > Self-Help > Personal Transformation", "Books > Health, Fitness & Dieting > Mental Health > Anxiety, Phobias & Panic Attacks"],
        "bisac": ["SEL023000 SELF-HELP / Personal Growth / General", "PSY022000 PSYCHOLOGY / Mental Health"],
    },
    "30-Day Challenge": {
        "categories": ["Books > Self-Help > Personal Transformation", "Books > Self-Help > Motivational"],
        "bisac": ["SEL023000 SELF-HELP / Personal Growth / General", "SEL027000 SELF-HELP / Motivational & Inspirational"],
    },
    "Planner": {
        "categories": ["Books > Arts & Photography > Graphic Design > Commercial > Calendars & Planners", "Books > Self-Help > Time Management"],
        "bisac": ["SEL043000 SELF-HELP / Time Management", "BUS107000 BUSINESS & ECONOMICS / Time Management"],
    },
    "Prompt Book": {
        "categories": ["Books > Self-Help > Journaling & Bullet Journaling", "Books > Crafts, Hobbies & Home > Games & Activities > Writing"],
        "bisac": ["SEL031000 SELF-HELP / Journaling", "GAM001000 GAMES & ACTIVITIES / General"],
    },
    "Activity Book": {
        "categories": ["Books > Children's Books > Activities, Crafts & Games > Activity Books", "Books > Arts & Photography > Drawing"],
        "bisac": ["JUV000000 JUVENILE FICTION / General", "GAM001000 GAMES & ACTIVITIES / General"],
    },
    "Self-help Guide": {
        "categories": ["Books > Self-Help > Personal Transformation", "Books > Health, Fitness & Dieting > Mental Health > Depression"],
        "bisac": ["SEL023000 SELF-HELP / Personal Growth / General", "SEL016000 SELF-HELP / Happiness"],
    },
    "default": {
        "categories": ["Books > Self-Help > Personal Transformation", "Books > Health, Fitness & Dieting"],
        "bisac": ["SEL023000 SELF-HELP / Personal Growth / General", "HEA000000 HEALTH & FITNESS / General"],
    },
}

def _get_kdp_categories(book_type: str) -> dict:
    for key in _KDP_CATEGORIES:
        if key.lower() in (book_type or "").lower():
            return _KDP_CATEGORIES[key]
    return _KDP_CATEGORIES["default"]


@app.post("/package", dependencies=[_AUTH])
async def generate_package(req: PackageRequest):
    tone_note = f"Tone/voice of the book: {req.tone}" if req.tone else ""
    persona_note = f"Target reader persona: {req.reader_persona}" if req.reader_persona else ""
    custom_note = f"\nCUSTOM RESTRICTIONS (MANDATORY — apply to every field):\n{req.custom_instructions}" if req.custom_instructions else ""
    stamp = now_stamp()
    _cats = _get_kdp_categories(req.book_type)
    _cat_hint = (
        f"\nREAL AMAZON CATEGORIES — use these exact browse paths (or close variants):\n"
        + "\n".join(f"  • {c}" for c in _cats["categories"])
        + f"\nREAL BISAC CODES — use these (adjust subject if needed):\n"
        + "\n".join(f"  • {b}" for b in _cats["bisac"])
    )

    # Determine target marketplace from language
    lang = getattr(req, 'language', 'English') or 'English'
    marketplace_map = {
        "Italian": "Amazon.it",
        "Dutch": "Amazon.nl",
        "German": "Amazon.de",
        "French": "Amazon.fr",
        "Spanish": "Amazon.es",
        "English": "Amazon.com",
    }
    marketplace = marketplace_map.get(lang, "Amazon.com")
    lang_note = f"Output language: {lang} (optimize for {marketplace})" if lang != "English" else "Output language: English (optimize for Amazon.com and Amazon.co.uk)"

    prompt = f"""You are an Amazon KDP publishing expert. Generate a complete KDP listing package.

Book: "{req.book_title}" — {req.book_subtitle}
Type: {req.book_type}
Trend: "{req.trend_name}"
Audience: {req.audience}
Platform: {req.trend_platform}
{tone_note}
{persona_note}
{lang_note}{custom_note}
Generated: {stamp}

IMPORTANT: Write title, subtitle, description, tagline, and keywords in {lang}.
Keywords must be search terms that {lang}-speaking readers actually use on {marketplace}.
{_cat_hint}

TITLE RULES — THIS IS CRITICAL, Amazon KDP REJECTS listings that violate this:
- "title": SHORT and punchy, under 60 characters. Must NOT contain a colon
  followed by a second long subtitle baked into it (e.g. do not write
  "Main Title: Long Secondary Subtitle Full Of Keywords" as the title)
- "subtitle": separate field for SEO/keywords, under 140 characters
- title + subtitle combined MUST be under 200 characters total — this is a
  hard Amazon limit and listings over it get rejected as "disappointing
  customer experience"

Return ONLY raw JSON. No markdown. ASCII-safe strings only.

{{"kdp":{{"title":"short punchy title, under 60 chars, no embedded subtitle","subtitle":"SEO subtitle, under 140 chars, title+subtitle combined under 200 chars total","pen_name":"believable author name fitting this niche and tone","pen_name_rationale":"1 sentence why this name works","description":"Full Amazon description 400-600 words. Use <b> for headers, <br> for breaks. Hook, benefits, who it is for.","short_description":"80-word mobile preview","keywords":["kw1","kw2","kw3","kw4","kw5","kw6","kw7"],"categories":["Primary Amazon category","Secondary Amazon category"],"bisac":["BISAC 1","BISAC 2"],"price_ebook":4.99,"price_paperback":12.99,"page_count_estimate":120,"trim_size":"6x9","tagline":"Punchy tagline under 15 words","canva_cover":{{"main_prompt":"60-80 word Canva AI image prompt for cover background. Mood colors lighting style. No text in scene.","style":"one-word style","color_palette":["#hex1","#hex2","#hex3"],"color_palette_names":["name1","name2","name3"],"font_title":"Canva font for title","font_subtitle":"Canva font for subtitle and author","layout_tip":"One sentence on placement","variation_1":"Alternative 40-word prompt","variation_2":"Alternative 40-word prompt","canva_steps":"4-5 step instructions for KDP-ready cover in Canva"}}}}}}"""

    try:
        text = await call_claude(prompt, 4000)
        result = parse_json_safe(text)
        kdp = result.get("kdp", {})
        title = kdp.get("title", "") or ""
        subtitle = kdp.get("subtitle", "") or ""
        warnings = []
        # Amazon KDP hard limit: title + subtitle combined <= 200 chars.
        # If Claude still baked a second subtitle into the title (the
        # "Title: Long Keyword-Stuffed Subtitle" pattern that gets listings
        # rejected as "disappointing customer experience"), split it off.
        if ":" in title and len(title) > 60:
            head, _, tail = title.partition(":")
            head, tail = head.strip(), tail.strip()
            if head and len(head) < 60:
                title = head
                subtitle = (tail + (" — " + subtitle if subtitle else "")).strip(" —")
                warnings.append("Titolo conteneva un sottotitolo incorporato — separato automaticamente")
        if len(title) + len(subtitle) > 200:
            max_subtitle = max(0, 200 - len(title))
            if len(subtitle) > max_subtitle:
                subtitle = subtitle[:max_subtitle].rsplit(" ", 1)[0]
                warnings.append("Sottotitolo troncato per rispettare il limite Amazon di 200 caratteri (titolo+sottotitolo)")
        kdp["title"] = title
        kdp["subtitle"] = subtitle
        if warnings:
            kdp["title_warnings"] = warnings
        result["kdp"] = kdp
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ══════════════════════════════════════════════════════════════
# REAL DOCX EXPORT
# ══════════════════════════════════════════════════════════════

@app.post("/export/docx", dependencies=[_AUTH])
async def export_docx(req: dict):
    """Generate a real .docx file from KDP book data + chapter content."""
    import io, re as _re
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    book = req.get("book", {})
    chapters_text = req.get("chapters_text", "")

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Title page ────────────────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run(book.get("title", "Untitled"))
    run.bold = True
    run.font.size = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if book.get("subtitle"):
        p2 = doc.add_paragraph()
        p2.add_run(book["subtitle"]).font.size = Pt(16)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if book.get("pen_name"):
        p3 = doc.add_paragraph()
        p3.add_run(f"by {book['pen_name']}").font.size = Pt(12)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── KDP Package ───────────────────────────────────────────
    doc.add_heading("KDP Package", level=1)

    def field(label, value):
        if not value:
            return
        p = doc.add_paragraph()
        p.add_run(label + ": ").bold = True
        p.add_run(str(value))

    field("Tagline",    book.get("tagline"))
    field("Pen Name",   (book.get("pen_name") or "") + (" — " + book["pen_name_rationale"] if book.get("pen_name_rationale") else ""))
    field("Prezzo",     f"eBook ${book.get('price_ebook','4.99')} | Paperback ${book.get('price_paperback','12.99')} | ~{book.get('page_count_estimate','120')} pagine | {book.get('trim_size','6x9')}")

    if book.get("description"):
        doc.add_heading("Descrizione Amazon", level=2)
        clean = _re.sub(r'<[^>]+>', ' ', book["description"]).strip()
        doc.add_paragraph(clean)

    if book.get("short_description"):
        doc.add_heading("Descrizione Breve (mobile)", level=2)
        doc.add_paragraph(book["short_description"])

    if book.get("keywords"):
        doc.add_heading("Keywords SEO", level=2)
        doc.add_paragraph(" | ".join(book["keywords"]))

    if book.get("categories"):
        doc.add_heading("Categorie Amazon", level=2)
        for cat in book["categories"]:
            doc.add_paragraph(f"• {cat}")

    if book.get("bisac"):
        field("BISAC", " | ".join(book["bisac"]))

    # ── Chapter content ───────────────────────────────────────
    if chapters_text:
        doc.add_page_break()
        doc.add_heading("Contenuto del Libro", level=1)
        for line in chapters_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            else:
                doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = _re.sub(r'[^a-zA-Z0-9_-]', '_', book.get("title", "book"))[:50] + ".docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
    )


# ══════════════════════════════════════════════════════════════
# APIFY INTEGRATION
# ══════════════════════════════════════════════════════════════

APIFY_TOKEN = env("APIFY_TOKEN")
APIFY_BASE = "https://api.apify.com/v2"

NICHE_CATEGORY_URLS = {
    "Self-help": "https://www.amazon.com/Best-Sellers-Books-Self-Help/zgbs/books/4736",
    "Self-help & Personal growth": "https://www.amazon.com/Best-Sellers-Books-Self-Help/zgbs/books/4736",
    "Health": "https://www.amazon.com/Best-Sellers-Books-Health-Fitness-Dieting/zgbs/books/6",
    "Health & Wellness": "https://www.amazon.com/Best-Sellers-Books-Health-Fitness-Dieting/zgbs/books/6",
    "Finance": "https://www.amazon.com/Best-Sellers-Books-Business-Money/zgbs/books/2",
    "Finance & Money": "https://www.amazon.com/Best-Sellers-Books-Business-Money/zgbs/books/2",
    "Relationships": "https://www.amazon.com/Best-Sellers-Books-Relationships-Parenting-Personal-Development/zgbs/books/48",
    "Relationships & Dating": "https://www.amazon.com/Best-Sellers-Books-Relationships-Parenting-Personal-Development/zgbs/books/48",
    "Productivity": "https://www.amazon.com/Best-Sellers-Books-Time-Management/zgbs/books/173514",
    "Mindset": "https://www.amazon.com/Best-Sellers-Books-Success-Self-Motivation/zgbs/books/4019",
    "Fitness": "https://www.amazon.com/Best-Sellers-Books-Exercise-Fitness/zgbs/books/31",
    "Parenting": "https://www.amazon.com/Best-Sellers-Books-Parenting-Relationships/zgbs/books/4735",
    "Business": "https://www.amazon.com/Best-Sellers-Books-Entrepreneurship/zgbs/books/12901",
    "Psychology": "https://www.amazon.com/Best-Sellers-Books-Psychology-Counseling/zgbs/books/25",
}

@app.get("/api/bestseller-links")
async def bestseller_links(niche: str = "", book_type: str = ""):
    """Return Amazon bestseller browse links for a niche + KDP category hints."""
    links = []
    # Match niche to category URL (fuzzy)
    niche_lower = niche.lower()
    for key, url in NICHE_CATEGORY_URLS.items():
        if any(w in niche_lower for w in key.lower().split()):
            links.append({"label": key, "url": url})
    if not links:
        links.append({"label": "All Books Bestsellers", "url": "https://www.amazon.com/Best-Sellers-Books/zgbs/books"})
    kdp_cats = _get_kdp_categories(book_type)
    return {"links": links[:3], "suggested_categories": kdp_cats["categories"], "suggested_bisac": kdp_cats["bisac"]}


@app.post("/api/similar-niches", dependencies=[_AUTH])
async def similar_niches(req: dict):
    """Return 4 adjacent KDP niches related to the selected one."""
    niche = (req.get("niche") or "").strip()
    if not niche:
        raise HTTPException(status_code=400, detail="niche required")
    prompt = (
        f'KDP niche research expert. Given the niche "{niche}", suggest exactly 4 adjacent/related niches that:\n'
        "- Appeal to a similar reader (overlapping pain points or interests)\n"
        "- Are distinct enough to be a separate book opportunity on Amazon\n"
        "- Have real search volume and KDP bestsellers\n\n"
        "Respond ONLY with a valid JSON array of 4 short niche names (2-5 words each):\n"
        '["Niche A", "Niche B", "Niche C", "Niche D"]'
    )
    try:
        text = await call_claude(prompt, max_tokens=200)
        parsed = parse_json_safe(text)
        if isinstance(parsed, list):
            return {"niches": [str(n) for n in parsed[:5]]}
        return {"niches": []}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/compare-outlines", dependencies=[_AUTH])
async def compare_outlines(req: dict):
    """Generate 3 parallel outline variants for side-by-side comparison."""
    trend = req.get("trend", "")
    book_type = req.get("book_type", "") or "self-help"
    title = req.get("title", "") or trend
    audience = req.get("audience", "general adult audience")
    custom_instructions = req.get("custom_instructions", "")
    _ci = f"\nCUSTOM RESTRICTIONS (mandatory): {custom_instructions}" if custom_instructions else ""

    styles = [
        ("🌊 Narrative Arc", "journey: Problem → Awareness → Transformation → Mastery — chapters build progressively on each other"),
        ("🔧 Modular/Reference", "self-contained modules — each chapter is complete and usable independently, tool-heavy"),
        ("⚡ 30-Day Challenge", "progressive daily/weekly actions — each chapter ends with a concrete exercise or challenge"),
    ]

    async def gen_one(style_name: str, style_desc: str) -> str:
        p = (
            f"Create a concise {book_type} book outline.\n"
            f'Title: "{title}"\n'
            f"Target audience: {audience}\n"
            f"Structure style: {style_desc}{_ci}\n\n"
            f"Format your response exactly like this:\n"
            f"## Style: {style_name}\n"
            "**Premise** (1 sentence): ...\n"
            "**Chapters:**\n"
            "1. Chapter Title — one-line description\n"
            "2. Chapter Title — one-line description\n"
            "(7-10 chapters total)\n"
            "**Conclusion:** one line\n\n"
            "Be specific and concrete. No placeholder text."
        )
        return await call_claude(p, max_tokens=800)

    try:
        results = await asyncio.gather(*[gen_one(s[0], s[1]) for s in styles])
        return {"variants": [{"style": styles[i][0], "outline": r} for i, r in enumerate(results)]}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


MARKET_GEO_MAP = {
    "English": "US",
    "Italiano": "IT",
    "Tedesco": "DE",
    "Spagnolo": "ES",
    "Francese": "FR",
    "Portoghese": "BR",
}


@app.post("/api/review-mining", dependencies=[_AUTH])
async def review_mining(req: dict):
    """Analyze pasted Amazon reviews → love/hate + differentiation opportunities."""
    reviews = (req.get("reviews") or "").strip()
    niche = req.get("niche", "")
    book_title = req.get("book_title", "")

    if len(reviews) < 80:
        raise HTTPException(status_code=400, detail="Incolla almeno 5-6 recensioni per un'analisi utile")

    title_note = f' (il mio libro: "{book_title}")' if book_title else ""
    prompt = (
        f'Amazon KDP market intelligence expert. Analyze these reader reviews for the "{niche}" niche{title_note}.\n\n'
        f"REVIEWS:\n{reviews[:4000]}\n\n"
        "Detect the language of the reviews and respond in that same language.\n"
        "Return ONLY valid JSON:\n"
        "{\n"
        '  "loves": ["3-5 things readers consistently PRAISE in the best books of this niche"],\n'
        '  "hates": ["3-5 recurring COMPLAINTS — what\'s missing or done poorly"],\n'
        '  "emotional_patterns": ["2-3 emotional phrases or words readers repeat — useful for copywriting"],\n'
        '  "differentiation_opportunities": ["3 concrete angles to stand out based on unmet needs"],\n'
        '  "ideal_book_description": "One sentence: what readers are ACTUALLY looking for, in their own words",\n'
        '  "warning_signals": ["1-2 red flags you must avoid or risk 1-star reviews"],\n'
        '  "sentiment_ratio": "X% positive / Y% negative / Z% mixed"\n'
        "}"
    )
    try:
        text = await call_claude(prompt, max_tokens=1500)
        return parse_json_safe(text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/price-optimizer", dependencies=[_AUTH])
async def price_optimizer(req: dict):
    """AI price recommendation for KDP ebook + paperback."""
    niche = req.get("niche", "")
    book_type = req.get("book_type", "")
    page_count = req.get("page_count", 120)
    competitor_prices = req.get("competitor_prices", [])
    current_price = req.get("current_price", 9.99)
    market = req.get("market", "US")

    comp_str = (
        f"Competitor paperback prices observed: {', '.join(f'${p}' for p in competitor_prices[:10])}"
        if competitor_prices else "No competitor prices provided."
    )
    prompt = (
        f"Amazon KDP pricing strategist. Recommend optimal prices for this book.\n\n"
        f"- Niche: {niche}\n"
        f"- Type: {book_type}\n"
        f"- Pages: {page_count}\n"
        f"- Market: Amazon {market}\n"
        f"- Current price considered: ${current_price}\n"
        f"- {comp_str}\n\n"
        "Consider KDP royalty tiers (35% under $2.99, 60% $2.99–$9.99), niche price sensitivity, "
        "competitive landscape, and value perception.\n\n"
        "Return ONLY valid JSON:\n"
        "{\n"
        '  "ebook_recommended": 4.99,\n'
        '  "ebook_rationale": "1 sentence why",\n'
        '  "paperback_recommended": 12.99,\n'
        '  "paperback_rationale": "1 sentence why",\n'
        '  "price_strategy": "penetration|competitive|premium",\n'
        '  "strategy_description": "2 sentences explaining the chosen strategy",\n'
        '  "launch_pricing": "Specific advice: start at $X for launch week, then move to $Y",\n'
        '  "risk_higher": "1 sentence: what happens if you price $1-2 more",\n'
        '  "risk_lower": "1 sentence: race-to-bottom risk or opportunity"\n'
        "}"
    )
    try:
        text = await call_claude(prompt, max_tokens=800)
        return parse_json_safe(text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/editor-pass", dependencies=[_AUTH])
async def editor_pass(req: dict):
    """AI editorial review of a content block for clarity, rhythm, style."""
    content = (req.get("content") or "").strip()
    label = req.get("label", "contenuto")
    tab = req.get("tab", "full")
    if len(content) < 100:
        raise HTTPException(status_code=400, detail="Contenuto troppo breve per una revisione utile")

    type_hints = {
        "outline": "book outline / chapter structure",
        "full": "full book chapter or section",
        "chapter": "individual chapter",
        "intro": "book introduction",
    }
    content_type = type_hints.get(tab, "book content")

    prompt = (
        f"You are a professional editor specializing in self-help and non-fiction books. "
        f"Review this {content_type} excerpt and provide actionable editorial suggestions.\n\n"
        f"SECTION: {label}\n\n"
        f"CONTENT:\n{content[:4000]}\n\n"
        "Analyze for: clarity of ideas, sentence rhythm and flow, repetitive words/phrases, "
        "weak verbs, passive voice overuse, structural issues, engagement level.\n\n"
        "IMPORTANT: Detect the language of the content and write ALL suggestions in that same language.\n\n"
        "Return ONLY valid JSON:\n"
        "{\n"
        '  "overall_score": 7,\n'
        '  "overall_notes": "2-3 sentence summary of the writing quality and main improvement area",\n'
        '  "suggestions": [\n'
        '    {\n'
        '      "issue": "Clarity|Rhythm|Repetition|Weak verb|Passive voice|Structure|Engagement",\n'
        '      "original": "exact phrase or sentence from text (max 15 words)",\n'
        '      "suggestion": "improved version",\n'
        '      "reason": "1 sentence why"\n'
        '    }\n'
        '  ]\n'
        "}\n"
        "Provide 3-6 specific, actionable suggestions. Reference exact phrases from the text."
    )
    try:
        text = await call_claude(prompt, max_tokens=1200)
        return parse_json_safe(text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/serie-planner", dependencies=[_AUTH])
async def serie_planner(req: dict):
    """Plan a 3-book series continuation from a KDP package."""
    title = req.get("title", "")
    niche = req.get("niche", "")
    book_type = req.get("book_type", "")
    description = (req.get("description") or "")[:800]

    prompt = (
        f"You are an Amazon KDP series strategist. Given the first book, plan 3 follow-up books "
        f"that expand the niche, deepen the audience relationship, and maximize reader LTV.\n\n"
        f"- Book 1 title: {title}\n"
        f"- Niche: {niche}\n"
        f"- Type: {book_type}\n"
        f"- Description: {description or 'N/A'}\n\n"
        "IMPORTANT: Detect the language from the title/description and write ALL output in that same language.\n\n"
        "Think: What is the natural 'next step' for a reader who loved Book 1? "
        "What deeper problem or adjacent goal can you address in Books 2, 3, 4?\n\n"
        "Return ONLY valid JSON:\n"
        "{\n"
        '  "series_name": "Series collection name",\n'
        '  "series_hook": "One-line series value proposition for readers",\n'
        '  "books": [\n'
        '    {\n'
        '      "number": 2,\n'
        '      "title": "Full title",\n'
        '      "subtitle": "Subtitle",\n'
        '      "angle": "Core angle/promise of this book in one sentence",\n'
        '      "why_buy": "Why a reader of Book 1 would want this next",\n'
        '      "differentiator": "What makes this distinct from Book 1"\n'
        '    }\n'
        '  ]\n'
        "}\n"
        "Plan exactly 3 follow-up books (numbers 2, 3, 4). Each must feel essential, not filler."
    )
    try:
        text = await call_claude(prompt, max_tokens=1500)
        return parse_json_safe(text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/translate-package", dependencies=[_AUTH])
async def translate_package(req: dict):
    """Translate and localize a KDP package to a target language/market."""
    language = req.get("language", "de")
    title = req.get("title", "")
    subtitle = req.get("subtitle", "")
    description = (req.get("description") or "")[:2000]
    keywords = req.get("keywords", [])

    lang_names = {
        "de": "German", "it": "Italian", "es": "Spanish",
        "fr": "French", "pt": "Portuguese", "nl": "Dutch",
    }
    lang_name = lang_names.get(language, language.upper())
    kw_str = ", ".join(keywords[:7]) if isinstance(keywords, list) else str(keywords)

    prompt = (
        f"You are a KDP localization expert for the Amazon {lang_name} marketplace.\n\n"
        f"Translate AND LOCALIZE (not just word-for-word translate) this KDP package:\n\n"
        f"TITLE: {title}\n"
        f"SUBTITLE: {subtitle}\n"
        f"DESCRIPTION:\n{description}\n"
        f"KEYWORDS: {kw_str}\n\n"
        f"Requirements:\n"
        f"1. Title/subtitle: natural {lang_name}, keep the marketing punch\n"
        f"2. Description: adapt idioms and tone for {lang_name} readers — culturally appropriate, not literal\n"
        f"3. Keywords: use actual search terms {lang_name} readers type on Amazon, NOT direct word translations\n"
        f"4. Note any key localization choices\n\n"
        "Return ONLY valid JSON:\n"
        "{\n"
        '  "title": "localized title",\n'
        '  "subtitle": "localized subtitle",\n'
        '  "description": "full localized description",\n'
        '  "keywords": ["kw1","kw2","kw3","kw4","kw5","kw6","kw7"],\n'
        '  "seo_notes": "2-3 sentences on localization choices and market-specific keywords"\n'
        "}"
    )
    try:
        text = await call_claude(prompt, max_tokens=2000)
        return parse_json_safe(text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/bsr-tracker", dependencies=[_AUTH])
async def bsr_tracker(req: dict):
    """Estimate BSR position and generate re-ranking suggestions for a published book."""
    title = req.get("title", "")
    niche = req.get("niche", "")
    asin = req.get("asin", "")
    review_count = int(req.get("review_count") or 0)
    price_ebook = float(req.get("price_ebook") or 4.99)
    days_since_pub = int(req.get("days_since_pub") or 30)
    keywords = req.get("keywords", [])
    kw_str = ", ".join(keywords[:7]) if isinstance(keywords, list) else ""

    prompt = (
        "You are an Amazon KDP ranking strategist. Estimate the BSR position and health of a book "
        "based on its metadata, then give specific re-ranking actions.\n\n"
        f"- ASIN: {asin}\n"
        f"- Title: {title}\n"
        f"- Niche: {niche}\n"
        f"- Reviews: {review_count}\n"
        f"- eBook price: ${price_ebook}\n"
        f"- Days since publication: {days_since_pub}\n"
        f"- Keywords: {kw_str or 'N/A'}\n\n"
        "BSR estimation logic: 0 reviews + <30 days → very poor (BSR 500k+); "
        "1-10 reviews → BSR 50k-500k; 11-50 reviews → BSR 5k-50k; "
        "50+ reviews → BSR <5k (if niche is healthy).\n\n"
        "IMPORTANT: Detect the language from the title/niche and write ALL text in that same language.\n\n"
        "Return ONLY valid JSON:\n"
        "{\n"
        '  "bsr_estimate": "#50,000 – #150,000 in Books",\n'
        '  "health_score": 4,\n'
        '  "competitive_position": "Weak / Building / Competitive / Strong",\n'
        '  "diagnosis": "2-sentence honest assessment of where the book stands",\n'
        '  "velocity_trend": "Declining|Stable|Growing",\n'
        '  "actions": [\n'
        '    "Specific action 1 — what to do and why",\n'
        '    "Specific action 2",\n'
        '    "Specific action 3",\n'
        '    "Specific action 4",\n'
        '    "Specific action 5"\n'
        '  ]\n'
        "}"
    )
    try:
        text = await call_claude(prompt, max_tokens=900)
        return parse_json_safe(text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/arc-sequence", dependencies=[_AUTH])
async def arc_sequence(req: dict):
    """Generate a 5-email ARC reviewer sequence for a KDP book launch."""
    title = req.get("title", "")
    subtitle = req.get("subtitle", "")
    pen_name = req.get("pen_name", "")
    niche = req.get("niche", "")
    description = (req.get("description") or "")[:600]
    launch_date = req.get("launch_date", "14 days from now")

    prompt = (
        "You are an email marketing expert for self-published authors. "
        "Write a 5-email sequence for ARC (Advance Review Copy) reviewers.\n\n"
        f"- Book title: {title}\n"
        f"- Subtitle: {subtitle}\n"
        f"- Author/pen name: {pen_name or 'the author'}\n"
        f"- Niche: {niche}\n"
        f"- Short description: {description or 'N/A'}\n"
        f"- Amazon launch date: {launch_date}\n\n"
        "Email schedule:\n"
        "  Email 1: Day 0 — Welcome + ARC delivery confirmation\n"
        "  Email 2: Day 3 — Reading check-in + value reminder\n"
        "  Email 3: Day 7 — Step-by-step how to leave a review\n"
        "  Email 4: Day 12 — Last chance reminder (2 days before launch)\n"
        "  Email 5: Launch day — Book is LIVE! Direct link + thank you\n\n"
        "IMPORTANT: Detect the language from the title/niche and write ALL emails in that same language. "
        "Make each email feel personal and warm, not corporate. Subject lines must have high open rates.\n\n"
        "Return ONLY valid JSON:\n"
        "{\n"
        '  "emails": [\n'
        '    {\n'
        '      "number": 1,\n'
        '      "timing": "Giorno 0 — Subito dopo l\'invio dell\'ARC",\n'
        '      "subject": "Email subject line",\n'
        '      "body": "Full email body with personalization placeholders like [NOME]"\n'
        '    }\n'
        '  ]\n'
        "}"
    )
    try:
        text = await call_claude(prompt, max_tokens=3500)
        return parse_json_safe(text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/niche-validator", dependencies=[_AUTH])
async def niche_validator(req: dict):
    """Full KDP niche profitability validation: demand, competition density, scores, go/no-go."""
    from urllib.parse import quote as url_quote
    niche = (req.get("niche") or "").strip()
    marketplace = req.get("marketplace", "us")
    if not niche:
        raise HTTPException(status_code=400, detail="Nicchia richiesta")

    # Step 1: Amazon keyword autocomplete — demand signal (fast, free)
    amazon_keywords: list[str] = []
    try:
        amazon_keywords = await _amazon_autocomplete(niche)
    except Exception:
        pass

    # Step 2: Apify Amazon book search — real BSR/review data (optional)
    books_data: list[dict] = []
    if APIFY_TOKEN:
        tld_map = {"us": "com", "de": "de", "it": "it", "es": "es", "fr": "fr"}
        tld = tld_map.get(marketplace, "com")
        search_url = (
            f"https://www.amazon.{tld}/s?k={url_quote(niche)}"
            f"&rh=n%3A283155&s=relevanceexprank"
        )
        try:
            items = await asyncio.wait_for(
                run_actor(
                    "junglee/amazon-crawler",
                    {
                        "categoryOrProductUrls": [{"url": search_url}],
                        "maxItemsPerStartUrl": 20,
                        "proxyConfiguration": {"useApifyProxy": True},
                    },
                ),
                timeout=100.0,
            )
            for item in items:
                if not isinstance(item, dict):
                    continue
                title = item.get("title") or item.get("name", "")
                if not title:
                    continue
                # BSR can be nested in various structures
                bsr_raw = (
                    item.get("bestsellersRank")
                    or item.get("bsr")
                    or item.get("bestSellersRank")
                )
                bsr = None
                if isinstance(bsr_raw, list) and bsr_raw:
                    first = bsr_raw[0]
                    bsr = first.get("rank") if isinstance(first, dict) else first
                elif isinstance(bsr_raw, (int, float)):
                    bsr = int(bsr_raw)
                books_data.append({
                    "title": str(title)[:80],
                    "bsr": bsr,
                    "reviews": item.get("reviewsCount") or item.get("reviews") or 0,
                    "price": item.get("price") or item.get("buyingPrice") or 0,
                    "rating": item.get("stars") or item.get("rating") or 0,
                })
        except Exception as e:
            print(f"[NicheValidator/Apify] {e}")

    # Step 3: Build context strings for Claude
    kw_str = (
        f"Amazon autocomplete keywords: {', '.join(amazon_keywords[:12])}"
        if amazon_keywords else ""
    )
    books_str = ""
    if books_data:
        lines = []
        for b in books_data[:15]:
            line = f"• {b['title']}"
            if b["bsr"]:
                line += f" | BSR #{b['bsr']:,}"
            if b["reviews"]:
                line += f" | {b['reviews']} reviews"
            if b["price"]:
                line += f" | ${b['price']}"
            lines.append(line)
        books_str = "\nTop books found on Amazon for this niche:\n" + "\n".join(lines)

    currency = "€" if marketplace in ("it", "de", "es", "fr") else "$"
    mkt_labels = {
        "us": "Amazon.com (US)", "de": "Amazon.de (Germany)",
        "it": "Amazon.it (Italy)", "es": "Amazon.es (Spain)", "fr": "Amazon.fr (France)",
    }
    mkt_label = mkt_labels.get(marketplace, "Amazon.com")
    mkt_domain = {"us": "Amazon.com", "de": "Amazon.de", "it": "Amazon.it", "es": "Amazon.es", "fr": "Amazon.fr"}.get(marketplace, "Amazon.com")

    prompt = (
        "You are a senior KDP niche analyst. Evaluate this niche for profitable self-publishing.\n\n"
        f"NICHE: {niche}\n"
        f"MARKETPLACE: {mkt_label}\n"
        f"CURRENCY: {currency}\n"
        f"{kw_str}\n"
        f"{books_str}\n\n"
        f"CRITICAL RULES:\n"
        f"- ALWAYS refer to the marketplace as '{mkt_domain}' — NEVER write 'Amazon.com' if the marketplace is not US\n"
        f"- Use '{currency}' for ALL prices, revenues, and estimates — NEVER use '$' if marketplace is not US\n"
        "- Detect the language from the niche keyword and write ALL text fields in that language.\n\n"
        "Scoring guide:\n"
        "- demand_score: reader search activity + purchase intent (1=nobody, 10=massive)\n"
        "- competition_score: market saturation (1=wide open, 10=impossible to enter)\n"
        "- entry_difficulty: reviews/quality threshold to rank on page 1 (1=easy 5 reviews, 10=need 500+)\n"
        "- profitability_score: overall opportunity = (demand - competition - difficulty) composite (1-10)\n\n"
        "Return ONLY valid JSON:\n"
        "{\n"
        '  "demand_score": 7,\n'
        '  "competition_score": 5,\n'
        '  "entry_difficulty": 4,\n'
        '  "profitability_score": 8,\n'
        '  "verdict": "GO",\n'
        '  "verdict_reason": "2 clear sentences on whether to enter this niche and exactly why",\n'
        '  "market_size": "Estimated ~X books, top-20 avg Y reviews",\n'
        f'  "revenue_estimate": "{currency}X-{currency}Y/month for a well-optimized top-20 book",\n'
        f'  "review_threshold": "~X reviews to appear on page 1",\n'
        f'  "best_angle": "The single most underserved sub-angle or reader segment",\n'
        f'  "ideal_price": "{currency}X.XX ebook / {currency}X.XX paperback",\n'
        '  "top_opportunities": ["specific opportunity 1", "opportunity 2", "opportunity 3"],\n'
        '  "key_risks": ["risk 1", "risk 2"],\n'
        '  "keyword_gems": ["high-potential keyword 1", "keyword 2", "keyword 3", "keyword 4"],\n'
        '  "trend_direction": "Rising",\n'
        '  "launch_urgency": "SOON",\n'
        '  "launch_window_weeks": 6,\n'
        '  "timing_advice": "1 sentence: when to publish and why"\n'
        "}\n"
        'verdict must be exactly one of: "GO", "CAUTIOUS", "NO-GO". '
        'trend_direction: "Rising"|"Stable"|"Declining"|"Seasonal". '
        'launch_urgency: "NOW" (publish within 2 weeks)|"SOON" (1-2 months)|"WAIT" (3-6 months, build quality)|"AVOID" (declining/saturated).'
    )
    try:
        text = await call_claude(prompt, max_tokens=1400)
        result = parse_json_safe(text)
        result["apify_data_used"] = bool(books_data)
        result["book_count_scraped"] = len(books_data)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/debug/env")
async def debug_env():
    """Show which env vars are present (keys only, no values) — for diagnosing Railway config."""
    import os
    apify_vars = {k: ("SET (non-empty)" if v else "SET (empty)") for k, v in os.environ.items() if "APIFY" in k.upper()}
    token_present = bool(APIFY_TOKEN)
    return {
        "APIFY_TOKEN_loaded": token_present,
        "apify_related_vars": apify_vars,
        "all_env_keys": sorted(os.environ.keys()),
    }


@app.post("/api/asin-reverse", dependencies=[_AUTH])
async def asin_reverse(req: dict):
    """Reverse-engineer a competitor ASIN: extract real niche, keywords, gaps, attack angle."""
    from urllib.parse import quote as url_quote
    asin = (req.get("asin") or "").strip().upper()
    marketplace = req.get("marketplace", "us")
    if not asin or len(asin) < 8:
        raise HTTPException(status_code=400, detail="ASIN non valido — deve essere tipo B0XXXXXXXX")

    # Step 1: Try Apify to get real product data
    tld_map = {"us": "com", "de": "de", "it": "it", "es": "es", "fr": "fr"}
    product_data: dict = {}
    apify_error_msg: str | None = None
    if APIFY_TOKEN:
        tld = tld_map.get(marketplace, "com")
        product_url = f"https://www.amazon.{tld}/dp/{asin}"
        try:
            items = await asyncio.wait_for(
                run_actor(
                    "junglee/amazon-crawler",
                    {
                        "categoryOrProductUrls": [{"url": product_url}],
                        "maxItemsPerStartUrl": 1,
                        "proxyConfiguration": {"useApifyProxy": True},
                    },
                ),
                timeout=90.0,
            )
            print(f"[AsinReverse] actor returned {len(items) if items else 0} items. Keys: {list(items[0].keys()) if items else []}")
            if items and isinstance(items[0], dict):
                it = items[0]
                bsr_raw = it.get("bestsellersRank") or it.get("bsr") or it.get("bestSellersRank")
                bsr = None
                if isinstance(bsr_raw, list) and bsr_raw:
                    first = bsr_raw[0]
                    bsr = first.get("rank") if isinstance(first, dict) else first
                elif isinstance(bsr_raw, (int, float)):
                    bsr = int(bsr_raw)
                product_data = {
                    "title": it.get("title") or it.get("name", ""),
                    "description": str(it.get("description") or it.get("about") or "")[:600],
                    "bsr": bsr,
                    "reviews": it.get("reviewsCount") or it.get("reviews") or 0,
                    "rating": it.get("stars") or it.get("rating") or 0,
                    "price": it.get("price") or it.get("buyingPrice") or 0,
                    "categories": it.get("breadCrumbs") or it.get("categories") or [],
                    "bullets": it.get("bullets") or [],
                }
        except Exception as e:
            apify_error_msg = str(e)[:300]
            print(f"[AsinReverse/Apify] {e}")

    context = f"ASIN: {asin}\nMarketplace: amazon.{tld_map.get(marketplace, 'com')}\n"
    if product_data.get("title"):
        context += f"Title: {product_data['title']}\n"
    if product_data.get("bsr"):
        context += f"BSR: #{product_data['bsr']:,}\n"
    if product_data.get("reviews"):
        context += f"Reviews: {product_data['reviews']} (rating: {product_data.get('rating',0)}★)\n"
    if product_data.get("price"):
        context += f"Price: ${product_data['price']}\n"
    if product_data.get("categories"):
        cats = [str(c) for c in product_data["categories"][:4]]
        context += f"Categories: {' > '.join(cats)}\n"
    if product_data.get("description"):
        context += f"Description: {product_data['description']}\n"
    # Without real product data Claude hallucinate — block early and return a clear error
    if not product_data:
        return {
            "real_niche": None,
            "primary_keywords": [],
            "positioning_angle": None,
            "reader_profile": None,
            "estimated_monthly_sales": None,
            "competitive_gap": None,
            "attack_angle": None,
            "title_formula": None,
            "strengths": [],
            "niche_for_validator": None,
            "apify_used": False,
            "error_no_data": True,
            "apify_error": apify_error_msg,
        }

    prompt = (
        "You are a KDP competitive intelligence analyst. Reverse-engineer this Amazon book "
        "to extract its niche strategy, keyword footprint, and competitive gaps.\n\n"
        f"{context}\n"
        "IMPORTANT: Detect the language from the title/description and write ALL text in that same language.\n\n"
        "Return ONLY valid JSON:\n"
        "{\n"
        '  "real_niche": "precise niche (more specific than genre — e.g. not \'self-help\' but \'mindfulness for burned-out moms\')",\n'
        '  "primary_keywords": ["kw1", "kw2", "kw3", "kw4", "kw5"],\n'
        '  "positioning_angle": "1 sentence: how this book frames its value proposition",\n'
        '  "reader_profile": "1 sentence: who exactly buys this (age, situation, pain point)",\n'
        '  "estimated_monthly_sales": "~X copies/month based on BSR",\n'
        '  "competitive_gap": "What this book fails to address that readers would still want",\n'
        '  "attack_angle": "Exactly how to position a competing book to win — specific angle, title direction, format difference",\n'
        '  "title_formula": "The underlying title formula (e.g. \'[Number] Ways to [Outcome] for [Audience]\')",\n'
        '  "strengths": ["what this book does well that you must match or beat"],\n'
        '  "use_as_niche": "clean niche keyword you can immediately use as a KDP niche input"\n'
        "}"
    )
    try:
        text = await call_claude(prompt, max_tokens=1000)
        result = parse_json_safe(text)
        result["asin"] = asin
        result["live_data"] = bool(product_data)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


async def run_actor(actor_id: str, input_data: dict, timeout_sec: int = 120) -> list:
    """Launch an Apify actor synchronously and return the dataset items."""
    if not APIFY_TOKEN:
        raise HTTPException(status_code=503, detail="APIFY_TOKEN non configurato — aggiungerlo nelle variabili d'ambiente Railway")
    actor_id_url = actor_id.replace("/", "~")
    async with httpx.AsyncClient(timeout=timeout_sec + 15) as client:
        try:
            run_res = await client.post(
                f"{APIFY_BASE}/acts/{actor_id_url}/runs",
                params={"token": APIFY_TOKEN, "waitForFinish": timeout_sec},
                json=input_data,
            )
            run_res.raise_for_status()
        except httpx.HTTPStatusError as e:
            raise HTTPException(
                status_code=e.response.status_code,
                detail=f"Apify actor '{actor_id}' HTTP {e.response.status_code} — {e.response.text[:300]}",
            )
        run_data = run_res.json().get("data", {})
        dataset_id = run_data.get("defaultDatasetId")
        if not dataset_id:
            raise HTTPException(status_code=502, detail="Apify actor non ha restituito un dataset ID")
        try:
            data_res = await client.get(
                f"{APIFY_BASE}/datasets/{dataset_id}/items",
                params={"token": APIFY_TOKEN, "format": "json"},
            )
            data_res.raise_for_status()
        except httpx.HTTPStatusError as e:
            raise HTTPException(
                status_code=e.response.status_code,
                detail=f"Apify dataset fetch HTTP {e.response.status_code}",
            )
        return data_res.json()


@app.post("/api/apify/trends", dependencies=[_AUTH])
async def apify_trends(req: dict):
    """Google Trends data for given keywords (automation-lab/google-trends-scraper)."""
    keywords = req.get("keywords", [])
    market_language = req.get("market_language", "English")
    geo = MARKET_GEO_MAP.get(market_language, "US")
    try:
        data = await run_actor(
            "automation-lab/google-trends-scraper",
            {"searchTerms": keywords, "timeRange": "today 3-m", "geo": geo, "outputMode": "interest_over_time"},
            timeout_sec=90,
        )
        return {"data": data, "geo": geo}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


_STOP = {
    # Articles, prepositions, conjunctions
    "the","a","an","or","and","but","for","with","from","by","to","of","in",
    "on","at","as","it","its","vs","vs.","&","i","my","your","our","their",
    "this","that","how","why","what","when","where","who","will","would",
    "can","could","do","does","did","not","no","so","than","then","if",
    "about","into","over","after","more","—","-","'s",
    # Be/have
    "is","are","was","were","be","been","being","have","has","had",
    # Common title adjectives — cause bad autocomplete (e.g. "first" → school bracelets)
    "first","second","third","new","little","big","great","good","best",
    "last","next","small","old","other","complete","ultimate","essential",
    "simple","easy","quick","fast","real","true","full","every","all","top",
    # Generic book-type words (too broad for autocomplete)
    "book","books","workbook","journal","notebook","planner","guide","manual",
    # Numbers as words
    "one","two","three","four","five","six","ten",
}

def _key_query(text: str, max_words: int = 4) -> str:
    """Extract the most meaningful words for an autocomplete query."""
    import re
    words = re.sub(r"[\"'()!?:;,.]", "", text).split()
    # Keep words: not in stop list AND at least 3 chars (skip 1-2 char tokens)
    key = [w for w in words if w.lower() not in _STOP and len(w) >= 3]
    short = " ".join(key[:max_words]) if key else " ".join(words[:max_words])
    return short or text[:40]


async def _amazon_autocomplete(query: str) -> list[str]:
    """Call Amazon's autocomplete API (2017 version) directly — no Apify, sub-second.
    Tries the full query first; falls back to key-words only if empty."""
    async def _fetch(prefix: str) -> list[str]:
        async with httpx.AsyncClient(timeout=8) as client:
            res = await client.get(
                "https://completion.amazon.com/api/2017/suggestions",
                params={
                    "lop": "en_US",
                    "site-variant": "desktop",
                    "category": "stripbooks",
                    "prefix": prefix,
                    "mid": "ATVPDKIKX0DER",
                    "alias": "stripbooks",
                },
                headers={
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
                    "Accept": "application/json",
                    "Referer": "https://www.amazon.com/",
                },
            )
            res.raise_for_status()
            payload = res.json()
            sugg = payload.get("suggestions", [])
            return [s["value"] for s in sugg if isinstance(s, dict) and s.get("value")]

    results = await _fetch(query)
    if not results:
        # Full title returned nothing — retry with extracted key words
        short = _key_query(query, max_words=3)
        if short and short.lower() != query.lower():
            results = await _fetch(short)
    if not results:
        # Last resort: first meaningful word only — skip if < 4 chars (avoids
        # language collisions like Italian "chi" → Amazon Chinese New Year results)
        first = _key_query(query, max_words=1)
        if first and len(first) >= 4 and first.lower() != query.lower():
            results = await _fetch(first)
    return results


@app.post("/api/apify/amazon-niche", dependencies=[_AUTH])
async def apify_amazon_niche(req: dict):
    """Amazon keyword autocomplete — direct call, no actor cold start."""
    keyword = req.get("keyword", "")
    if not keyword:
        return {"data": []}
    try:
        suggestions = await _amazon_autocomplete(keyword)
        return {"data": [{"platform": "amazon", "suggestions": suggestions}]}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/apify/amazon-best", dependencies=[_AUTH])
async def apify_amazon_best(req: dict):
    """Amazon keyword autocomplete for a niche — direct call, no actor cold start."""
    niche = req.get("niche", "")
    if not niche:
        return {"data": []}
    try:
        suggestions = await _amazon_autocomplete(niche)
        return {"data": [{"platform": "amazon", "suggestions": suggestions}]}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def _shorten_query(q: str, max_words: int = 4) -> str:
    """Strip punctuation/stop-words and keep the most meaningful words."""
    import re
    stop = {"la","il","lo","le","i","gli","un","una","del","della","di","da","per","con","su","e","a"}
    words = [w for w in re.split(r'[\s\-–—]+', q) if w and w.lower() not in stop]
    return " ".join(words[:max_words])


async def _autocomplete_google(query: str, client_param: str = "firefox", lang: str = "it") -> list[str]:
    short = _shorten_query(query)
    for q in ([query, short] if short != query else [query]):
        for url in [
            "https://suggestqueries.google.com/complete/search",
            "https://clients1.google.com/complete/search",
        ]:
            try:
                params = {"client": client_param, "q": q, "hl": lang}
                async with httpx.AsyncClient(timeout=8) as c:
                    r = await c.get(url, params=params, headers={"User-Agent": "Mozilla/5.0"})
                    data = r.json()
                    results = data[1] if isinstance(data, list) and len(data) > 1 else []
                    if results:
                        return [str(x) for x in results][:15]
            except Exception:
                continue
    return []


async def _autocomplete_pinterest(query: str) -> list[str]:
    short = _shorten_query(query)
    for q in ([short, query] if short != query else [query]):
        try:
            url = "https://www.pinterest.com/typeahead/search/"
            params = {"q": q, "rs": "ac", "prefix": "/search/"}
            async with httpx.AsyncClient(timeout=10, follow_redirects=True) as c:
                r = await c.get(url, params=params,
                                headers={"User-Agent": "Mozilla/5.0", "Accept": "application/json"})
                data = r.json()
                items = data.get("resource_response", {}).get("data", []) or []
                result = [i.get("display") or i.get("term", "") for i in items
                          if i.get("display") or i.get("term")][:15]
                if result:
                    return result
        except Exception:
            continue
    return []


async def _autocomplete_tiktok(query: str) -> list[str]:
    short = _shorten_query(query)
    for q in ([short, query] if short != query else [query]):
        try:
            url = "https://www.tiktok.com/api/search/general/preview/"
            params = {"keyword": q, "from_page": "fyp"}
            async with httpx.AsyncClient(timeout=10) as c:
                r = await c.get(url, params=params,
                                headers={"User-Agent": "Mozilla/5.0", "Referer": "https://www.tiktok.com/"})
                data = r.json()
                items = data.get("sug_list") or data.get("data", []) or []
                if isinstance(items, list) and items and isinstance(items[0], dict):
                    result = [i.get("keyword") or i.get("word", "")
                              for i in items if i.get("keyword") or i.get("word")][:15]
                    if result:
                        return result
        except Exception:
            continue
    # Fallback: YouTube suggestions
    return await _autocomplete_google(short or query, client_param="youtube")


@app.post("/api/apify/keywords", dependencies=[_AUTH])
async def apify_keywords(req: dict):
    """Multi-platform keyword suggestions via direct autocomplete APIs (no Apify actor)."""
    query = req.get("query", "")
    platforms = req.get("platforms", ["google", "amazon", "pinterest"])
    longtail = req.get("longtail", False)

    if not query:
        return {"data": [], "longtail": []}

    platform_set = set(platforms)

    async def fetch_platform(plat: str):
        try:
            if plat == "google":
                sugg = await _autocomplete_google(query)
            elif plat == "amazon":
                sugg = await _amazon_autocomplete(query)
            elif plat == "pinterest":
                sugg = await _autocomplete_pinterest(query)
            elif plat == "tiktok":
                sugg = await _autocomplete_tiktok(query)
            elif plat == "youtube":
                sugg = await _autocomplete_google(query, client_param="youtube")
            elif plat == "instagram":
                # No public autocomplete — reuse Google with #hashtag framing
                sugg = await _autocomplete_google(f"{query} instagram")
            else:
                sugg = []
            return {"platform": plat, "suggestions": [s for s in sugg if s][:15]}
        except Exception:
            return {"platform": plat, "suggestions": []}

    tasks = [fetch_platform(p) for p in platforms if p in platform_set]
    results = await asyncio.gather(*tasks, return_exceptions=True)
    sugg_data = [r for r in results if isinstance(r, dict)]

    # Fill empty platforms with Claude-generated suggestions
    empty_platforms = [r["platform"] for r in sugg_data if not r.get("suggestions")]
    if empty_platforms:
        short = _shorten_query(query)
        fill_prompt = f"""You are a keyword research expert.
Topic: "{short or query}"
Generate 10 realistic search suggestions a user would type on each of these platforms: {', '.join(empty_platforms)}.
Suggestions should reflect how real users search on that platform (TikTok = short/trending, Instagram = hashtag-style, Pinterest = visual/tutorial, YouTube = "how to", Google = questions/intent).
Return ONLY valid JSON: {{"platform_name": ["suggestion1", "suggestion2", ...]}}"""
        try:
            raw = await call_claude(fill_prompt, 600)
            if raw.startswith("```"):
                raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
            filled = json.loads(raw)
            for item in sugg_data:
                if not item.get("suggestions") and item["platform"] in filled:
                    item["suggestions"] = [s for s in filled[item["platform"]] if s][:12]
        except Exception:
            pass

    # Long-tail: Claude-generated variations (no Apify actor needed)
    lt_data = []
    if longtail:
        all_kws = [s for r in sugg_data for s in r.get("suggestions", [])][:20]
        lt_prompt = f"""Given these keyword suggestions for "{query}": {', '.join(all_kws[:15])}

Return a JSON array of 15 long-tail keyword objects with estimated metrics:
[{{"keyword":"...","searchVolume":1200,"cpc":0.45,"competition":0.3}}]
Use realistic ranges for book/publishing niche: volume 100-5000, cpc $0.20-$2.00, competition 0.1-0.9.
Return ONLY valid JSON array."""
        try:
            raw = await call_claude(lt_prompt, 800)
            if raw.startswith("```"):
                raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
            lt_data = json.loads(raw)
            if not isinstance(lt_data, list):
                lt_data = []
        except Exception:
            lt_data = []

    return {"data": sugg_data, "longtail": lt_data}


@app.post("/api/apify/captions", dependencies=[_AUTH])
async def apify_captions(req: dict):
    """Generate social media captions for all platforms using Claude."""
    platforms = req.get("platforms", ["instagram", "facebook"])
    tone = req.get("tone", "warm")
    language = req.get("language", "it")
    custom_context = req.get("customContext", "")
    video_url = req.get("videoUrl", "")

    lang_label = {"it": "Italian", "en": "English", "it+en": "Italian with some English phrases"}.get(language, "Italian")
    tone_label = {"warm": "warm and friendly", "professional": "professional", "funny": "fun and playful", "urgent": "urgent and compelling"}.get(tone, tone)

    # Fetch real book info from Amazon URL so captions are accurate
    book_info = await fetch_amazon_book_info(video_url) if video_url else {}
    book_ctx_parts = []
    if book_info.get("title"):
        book_ctx_parts.append(f'Book title: "{book_info["title"]}"')
    if book_info.get("author"):
        book_ctx_parts.append(f'Author: {book_info["author"]}')
    if book_info.get("description"):
        book_ctx_parts.append(f'Description: {book_info["description"]}')
    if video_url:
        book_ctx_parts.append(f'Amazon link: {video_url}')
    book_ctx = "\n".join(book_ctx_parts) if book_ctx_parts else (custom_context or "KDP book on Amazon")

    platforms_str = ", ".join(platforms)
    prompt = f"""You are a social media expert specializing in book promotion for Amazon KDP.

Book info:
{book_ctx}
{f"Additional context: {custom_context}" if custom_context and book_ctx_parts else ""}

Generate ONE promotional post for EACH of these platforms: {platforms_str}

Rules:
- Language: {lang_label}
- Tone: {tone_label}
- Each post must be platform-native (length, style, emoji usage)
- Instagram: 150-220 chars + 8-10 hashtags
- Facebook: 180-280 chars, conversational, 3-5 hashtags
- TikTok: short punchy hook + 5-7 trending hashtags
- Pinterest: descriptive, keyword-rich, 100-150 chars + 5 hashtags
- Twitter: max 240 chars, punchy, 2-3 hashtags
- LinkedIn: professional, 200-300 chars, no hashtags or 2-3 professional ones

Return ONLY valid JSON array, no markdown, no extra text:
[
  {{"platform":"instagram","caption":"...","hashtags":["tag1","tag2"]}},
  {{"platform":"facebook","caption":"...","hashtags":["tag1"]}},
  ...one object per platform...
]"""

    try:
        raw = await call_claude(prompt, 1500)
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
        items = json.loads(raw)
        if not isinstance(items, list):
            items = []
        return {"data": items, "book_detected": book_info.get("title") or None}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/ads/campaign", dependencies=[_AUTH])
async def generate_ads_campaign(req: dict):
    """Generate Amazon Sponsored Products campaign structure using Claude, return JSON + CSV."""
    book_title = req.get("bookTitle", "")
    niche = req.get("niche", "")
    keywords = req.get("keywords", [])
    daily_budget = float(req.get("dailyBudget", 5.0))
    marketplace = req.get("marketplace", "US")
    asin = req.get("asin", "")

    kw_block = "\n".join(f"- {kw}" for kw in keywords[:30]) if keywords else "(generate relevant keywords for this niche)"

    prompt = f"""You are an Amazon Advertising expert specializing in KDP book campaigns.

Book: "{book_title}"
Niche: {niche}
Marketplace: Amazon {marketplace}
Daily budget: ${daily_budget:.2f}
{f'ASIN: {asin}' if asin else ''}
Seed keywords:
{kw_block}

Generate a Sponsored Products (manual targeting) campaign with 3 ad groups:
- Exact Match: 6-8 high-intent keywords, highest bids
- Phrase Match: 8-12 keywords, medium bids
- Broad Match: 10-15 broad keywords, lowest bids

Books typically bid $0.25-$1.50 CPC. Scale bids to niche competitiveness.
Also produce 12-15 negative exact keywords to block irrelevant clicks.

Return ONLY valid JSON (no markdown, no extra text):
{{
  "campaign_name": "SP - {book_title[:40]} - Manual",
  "daily_budget": {daily_budget},
  "targeting_type": "MANUAL",
  "bidding_strategy": "DYNAMIC_BIDS_DOWN_ONLY",
  "strategy_note": "2-3 sentence rationale",
  "ad_groups": [
    {{
      "name": "Exact Match",
      "default_bid": 0.75,
      "keywords": [{{"keyword": "...", "match_type": "EXACT", "bid": 0.80}}]
    }},
    {{
      "name": "Phrase Match",
      "default_bid": 0.50,
      "keywords": [{{"keyword": "...", "match_type": "PHRASE", "bid": 0.55}}]
    }},
    {{
      "name": "Broad Match",
      "default_bid": 0.30,
      "keywords": [{{"keyword": "...", "match_type": "BROAD", "bid": 0.35}}]
    }}
  ],
  "negative_keywords": ["keyword1", "keyword2"]
}}"""

    try:
        raw = await call_claude(prompt, 2500)
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
        campaign = json.loads(raw)

        # Build Amazon Sponsored Products bulk-upload CSV
        import csv, io
        from datetime import datetime, timedelta
        start_date = datetime.utcnow().strftime("%Y%m%d")
        cols = [
            "Product","Entity","Operation","Campaign ID","Ad Group ID","Portfolio ID",
            "Ad ID","Keyword ID","Product Targeting ID","Campaign Name","Ad Group Name",
            "Start Date","End Date","Targeting Type","State","Daily Budget","SKU","ASIN",
            "Ad Group Default Bid","Bid","Keyword Text","Match Type","Bidding Strategy",
            "Placement","Percentage","Product Targeting Expression"
        ]
        rows = []

        def row(**kw):
            r = {c: "" for c in cols}
            r.update(kw)
            return r

        cname = campaign.get("campaign_name", f"SP - {book_title}")

        # Campaign row
        rows.append(row(
            Product="Sponsored Products",
            Entity="Campaign",
            Operation="create",
            **{"Campaign Name": cname,
               "Targeting Type": "MANUAL",
               "State": "enabled",
               "Daily Budget": f"{campaign.get('daily_budget', daily_budget):.2f}",
               "Start Date": start_date,
               "Bidding Strategy": campaign.get("bidding_strategy", "DYNAMIC_BIDS_DOWN_ONLY")}
        ))

        for ag in campaign.get("ad_groups", []):
            ag_name = ag.get("name", "Ad Group")
            # Ad group row
            rows.append(row(
                Product="Sponsored Products",
                Entity="Ad Group",
                Operation="create",
                **{"Campaign Name": cname,
                   "Ad Group Name": ag_name,
                   "Ad Group Default Bid": f"{ag.get('default_bid', 0.50):.2f}",
                   "State": "enabled"}
            ))
            # Product ad row (ASIN if provided)
            if asin:
                rows.append(row(
                    Product="Sponsored Products",
                    Entity="Product Ad",
                    Operation="create",
                    **{"Campaign Name": cname,
                       "Ad Group Name": ag_name,
                       "ASIN": asin,
                       "State": "enabled"}
                ))
            # Keyword rows
            for kw in ag.get("keywords", []):
                rows.append(row(
                    Product="Sponsored Products",
                    Entity="Keyword",
                    Operation="create",
                    **{"Campaign Name": cname,
                       "Ad Group Name": ag_name,
                       "Keyword Text": kw.get("keyword", ""),
                       "Match Type": kw.get("match_type", "EXACT").upper(),
                       "Bid": f"{kw.get('bid', ag.get('default_bid', 0.50)):.2f}",
                       "State": "enabled"}
                ))

        # Negative keywords (campaign-level, exact)
        for neg in campaign.get("negative_keywords", []):
            rows.append(row(
                Product="Sponsored Products",
                Entity="Negative keyword",
                Operation="create",
                **{"Campaign Name": cname,
                   "Keyword Text": neg,
                   "Match Type": "NEGATIVE EXACT",
                   "State": "enabled"}
            ))

        buf = io.StringIO()
        writer = csv.DictWriter(buf, fieldnames=cols)
        writer.writeheader()
        writer.writerows(rows)
        csv_content = buf.getvalue()

        return {"data": campaign, "csv": csv_content}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/apify/video-story", dependencies=[_AUTH])
async def apify_video_story(req: dict):
    """Generate a promotional video script + storyboard using Claude."""
    prompt = req.get("prompt", "")
    duration = req.get("duration", 30)
    output_platform = req.get("outputPlatform", "instagram")
    book_title = req.get("bookTitle", "")
    niche = req.get("niche", "")

    context = prompt or f'Book "{book_title}" — niche: {niche}'
    sys_prompt = (
        f"You are a professional social media video director specializing in book promotion. "
        f"Create a {duration}-second video script for {output_platform}. "
        f"Output ONLY valid JSON, no markdown, no extra text."
    )
    user_prompt = (
        f"Create a {duration}s promotional video script for: {context}\n\n"
        f"Return JSON with this exact structure:\n"
        f'{{"title":"...", "hook":"...(0-3s hook text)", '
        f'"scenes":[{{"time":"0-5s","visual":"scene description","voiceover":"spoken text"}}, ...], '
        f'"cta":"...call to action text", "music":"...suggested music genre/mood", '
        f'"caption":"...ready-to-post social caption with hashtags"}}'
    )
    try:
        msg = claude.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=1200,
            system=sys_prompt,
            messages=[{"role": "user", "content": user_prompt}],
        )
        raw = msg.content[0].text.strip()
        # Strip markdown code fences if present
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
        script = json.loads(raw)
        return {"script": script}
    except json.JSONDecodeError:
        return {"script": {"title": context, "raw": raw if 'raw' in dir() else ""}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/apify/video-ugc", dependencies=[_AUTH])
async def apify_video_ugc(req: dict):
    """Generate a UGC-style video script from a cover image using Claude."""
    image_url = req.get("imageUrl", "")
    platform = req.get("platform", "tiktok")
    duration = req.get("duration", 30)
    book_title = req.get("bookTitle", "")
    niche = req.get("niche", "")

    context = f'Book: "{book_title}", niche: {niche}' if book_title else "KDP book"
    sys_prompt = (
        f"You are a UGC content creator specializing in book promotion on {platform}. "
        f"Create an authentic {duration}-second UGC video script. "
        f"Output ONLY valid JSON, no markdown, no extra text."
    )
    cover_note = f" Cover image: {image_url}" if image_url else ""
    user_prompt = (
        f"Create a {duration}s UGC-style video script for {platform}. {context}.{cover_note}\n\n"
        f"Return JSON with this exact structure:\n"
        f'{{"title":"...", "hook":"...grabby opening line", '
        f'"scenes":[{{"time":"0-5s","action":"what creator does on camera","text":"spoken/overlay text"}}, ...], '
        f'"cta":"...end CTA", "music":"...trending audio suggestion", '
        f'"caption":"...caption with hashtags for {platform}"}}'
    )
    try:
        msg = claude.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=1200,
            system=sys_prompt,
            messages=[{"role": "user", "content": user_prompt}],
        )
        raw = msg.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
        script = json.loads(raw)
        return {"script": script}
    except json.JSONDecodeError:
        return {"script": {"title": context, "raw": raw if 'raw' in dir() else ""}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/apify/social-post", dependencies=[_AUTH])
async def apify_social_post(req: dict):
    """Legacy endpoint — kept for backward compat, returns Pinterest link only."""
    caption = req.get("caption", "")
    platforms = req.get("platforms", [])
    results = []
    for platform in platforms:
        if platform == "pinterest":
            results.append({"platform": "pinterest", "success": True,
                            "message": "Caption pronta — copia il testo e incollalo su Pinterest",
                            "postUrl": "https://www.pinterest.com/pin-builder/", "copyText": caption})
        else:
            results.append({"platform": platform, "success": False,
                            "message": f"Usa il nuovo tab Blotato per pubblicare su {platform.capitalize()}."})
    return {"data": results}


_BLOTATO_BASE = "https://backend.blotato.com"


@app.post("/api/blotato/accounts", dependencies=[_AUTH])
async def blotato_accounts(req: dict):
    """Return the Blotato user's connected social accounts."""
    api_key = req.get("apiKey", "").strip()
    if not api_key:
        raise HTTPException(status_code=400, detail="Blotato API key mancante")
    async with httpx.AsyncClient(timeout=20) as client:
        r = await client.get(
            f"{_BLOTATO_BASE}/v2/users/me/accounts",
            headers={"Authorization": f"Bearer {api_key}"},
        )
        print(f"[Blotato accounts] status={r.status_code} body={r.text[:500]}")
        if r.status_code == 401:
            raise HTTPException(status_code=401, detail="API key Blotato non valida — controlla su my.blotato.com/settings/api")
        if r.status_code != 200:
            raise HTTPException(status_code=r.status_code, detail=r.text[:300])
        body = r.json()
        # Normalise: Blotato may return {accounts:[...]}, {data:[...]}, or [...] directly
        if isinstance(body, list):
            return {"accounts": body, "_raw": body}
        if isinstance(body, dict):
            accounts = body.get("accounts") or body.get("data") or body.get("items") or body.get("socialAccounts") or []
            return {"accounts": accounts, "_raw": body}
        return {"accounts": [], "_raw": body}


@app.post("/api/blotato/post", dependencies=[_AUTH])
async def blotato_post(req: dict):
    """Publish one or more posts via Blotato (one call per account)."""
    api_key = req.get("apiKey", "").strip()
    if not api_key:
        raise HTTPException(status_code=400, detail="Blotato API key mancante")
    posts = req.get("posts", [])
    if not posts:
        raise HTTPException(status_code=400, detail="Nessun post specificato")

    _RESERVED = {"apiKey"}
    results = []
    async with httpx.AsyncClient(timeout=60) as client:
        for post in posts:
            platform   = post.get("platform", "")
            media_urls = post.get("mediaUrls", [])
            content: dict = {"platform": platform, "text": post.get("text", ""), "mediaUrls": media_urls}
            # targetType mirrors platform name (Blotato enum)
            target: dict = {"targetType": platform}
            if post.get("boardId"):      target["boardId"]      = post["boardId"]
            if post.get("pageId"):       target["pageId"]       = post["pageId"]
            if post.get("privacyLevel"): target["privacyLevel"] = post["privacyLevel"]
            if post.get("privacyStatus"):target["privacyStatus"]= post["privacyStatus"]
            if post.get("title"):        content["title"]       = post["title"]
            post_obj: dict = {
                "accountId": post.get("accountId", ""),
                "content": content,
                "target": target,
            }
            if post.get("scheduledTime"):    post_obj["scheduledTime"]    = post["scheduledTime"]
            if post.get("useNextFreeSlot"):  post_obj["useNextFreeSlot"]  = post["useNextFreeSlot"]
            payload = {"post": post_obj}
            try:
                r = await client.post(
                    f"{_BLOTATO_BASE}/v2/posts",
                    headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                    json=payload,
                )
                if r.status_code in (200, 201):
                    body = r.json() if r.content else {}
                    results.append({"success": True, "platform": post.get("platform"), **body})
                else:
                    results.append({"success": False, "platform": post.get("platform"), "error": r.text[:300]})
            except Exception as e:
                results.append({"success": False, "platform": post.get("platform"), "error": str(e)})

    return {"results": results}


@app.post("/api/blotato/pinterest-boards", dependencies=[_AUTH])
async def blotato_pinterest_boards(req: dict):
    """Return Pinterest boards for a given Blotato accountId — tries multiple endpoints."""
    api_key = req.get("apiKey", "").strip()
    account_id = req.get("accountId", "").strip()
    if not api_key:
        raise HTTPException(status_code=400, detail="Blotato API key mancante")
    if not account_id:
        raise HTTPException(status_code=400, detail="accountId mancante")
    headers = {"Authorization": f"Bearer {api_key}"}
    # Blotato uses blotato-api-key header + specific path for Pinterest boards
    headers_with_key = {"blotato-api-key": api_key, "Authorization": f"Bearer {api_key}"}
    candidates = [
        (f"{_BLOTATO_BASE}/v2/social/pinterest/boards?accountId={account_id}", headers_with_key),
        (f"{_BLOTATO_BASE}/v2/social/pinterest/boards?accountId={account_id}", headers),
        (f"{_BLOTATO_BASE}/v2/users/me/accounts/{account_id}/pinterest-boards", headers),
        (f"{_BLOTATO_BASE}/v2/users/me/accounts/{account_id}/boards", headers),
    ]
    async with httpx.AsyncClient(timeout=20) as client:
        last_status, last_body = 404, ""
        for url, hdrs in candidates:
            r = await client.get(url, headers=hdrs)
            print(f"[Pinterest boards] {url} → {r.status_code} {r.text[:200]}")
            if r.status_code == 200:
                body = r.json()
                boards = body if isinstance(body, list) else (
                    body.get("boards") or body.get("data") or body.get("items") or []
                )
                return {"boards": boards, "_raw": body}
            last_status, last_body = r.status_code, r.text
        raise HTTPException(status_code=last_status, detail=f"Impossibile caricare le bacheche Pinterest: {last_body[:300]}")


@app.post("/api/apify/influencers", dependencies=[_AUTH])
async def apify_influencers(req: dict):
    """Find influencers by niche (easyapi/find-my-influencers)."""
    niche = req.get("niche", "")
    platforms = req.get("platforms", ["instagram", "tiktok"])
    try:
        data = await run_actor(
            "easyapi/find-my-influencers",
            {"niche": niche, "platforms": platforms, "limit": 20},
            timeout_sec=120,
        )
        return {"data": data}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ══════════════════════════════════════════════════════════════
# CUSTOMER AVATAR GENERATOR
# ══════════════════════════════════════════════════════════════

class AvatarRequest(BaseModel):
    reviews: str
    book_title: Optional[str] = ""
    niche: Optional[str] = ""
    audience: Optional[str] = ""

@app.post("/api/amazon-reviews", dependencies=[_AUTH])
async def fetch_amazon_reviews(req: dict):
    """Fetch real Amazon reviews for one or more ASINs via Apify — feeds Customer Avatar Generator."""
    # Support both single asin and asins array
    raw_asins = req.get("asins") or ([req.get("asin")] if req.get("asin") else [])
    asins = [a.strip().upper() for a in raw_asins if a and len(a.strip()) >= 8]
    marketplace = req.get("marketplace", "us")
    max_per_asin = min(int(req.get("max_reviews", 50)), 100)

    if not asins:
        raise HTTPException(status_code=400, detail="Nessun ASIN valido fornito")
    if not APIFY_TOKEN:
        raise HTTPException(status_code=503, detail="APIFY_TOKEN non configurato")

    tld_map = {"us": "com", "de": "de", "it": "it", "es": "es", "fr": "fr", "uk": "co.uk"}
    tld = tld_map.get(marketplace, "com")

    async def fetch_for_asin(asin: str) -> list:
        actor_configs = [
            ("automation-lab/amazon-reviews-scraper", {
                "asins": [asin],
                "maxReviews": max_per_asin,
                "reviewsCount": max_per_asin,
            }),
            ("epctex/amazon-reviews-scraper", {
                "startUrls": [{"url": f"https://www.amazon.{tld}/dp/{asin}"}],
                "maxItems": max_per_asin,
            }),
        ]
        for actor_id, actor_input in actor_configs:
            try:
                result = await asyncio.wait_for(run_actor(actor_id, actor_input), timeout=90.0)
                if result:
                    print(f"[AmazonReviews] {asin}: {actor_id} → {len(result)} items, keys={list(result[0].keys()) if result else []}")
                    return result
            except Exception as e:
                print(f"[AmazonReviews] {asin}: {actor_id} failed: {e}")
        return []

    # Fetch all ASINs concurrently
    all_results = await asyncio.gather(*[fetch_for_asin(a) for a in asins])
    all_items = [item for sub in all_results for item in sub]

    if not all_items:
        raise HTTPException(status_code=404, detail=f"Nessuna recensione trovata per: {', '.join(asins)}")

    # Debug: log actual keys so we can fix field mapping
    lines = []
    for it in all_items:
        if not isinstance(it, dict):
            continue
        rating = it.get("rating") or it.get("stars") or it.get("ratingScore") or ""
        title = (it.get("title") or it.get("reviewTitle") or it.get("reviewHeadline")
                 or it.get("headline") or "")
        text = (it.get("text") or it.get("reviewText") or it.get("body")
                or it.get("reviewBody") or it.get("content") or it.get("review")
                or it.get("reviewContent") or it.get("description") or "")
        if not text:
            continue
        try:
            star = "★" * int(float(rating)) if rating and 1 <= float(str(rating)) <= 5 else ""
        except (ValueError, TypeError):
            star = ""
        lines.append(f"{star} {title}\n{text}".strip())

    if not lines:
        sample_keys = list(all_items[0].keys()) if all_items and isinstance(all_items[0], dict) else []
        raise HTTPException(
            status_code=404,
            detail=f"Nessun testo recensione estratto. Campi disponibili: {sample_keys}"
        )

    return {
        "reviews_text": "\n\n".join(lines),
        "count": len(lines),
        "asins": asins,
        "marketplace": marketplace,
    }


@app.post("/api/avatar", dependencies=[_AUTH])
async def generate_avatar(req: AvatarRequest):
    """Generate 2 customer avatars from competitor reviews using Claude."""
    context_parts = []
    if req.book_title:
        context_parts.append(f"Book: {req.book_title}")
    if req.niche:
        context_parts.append(f"Niche: {req.niche}")
    if req.audience:
        context_parts.append(f"Audience hint: {req.audience}")
    context_line = " | ".join(context_parts) if context_parts else "KDP non-fiction book"

    system_prompt = (
        "You are an expert customer research analyst for Amazon KDP publishers. "
        "Analyze competitor reviews to identify distinct buyer personas. "
        "IMPORTANT: Detect the language of the reviews and input data, then respond entirely "
        "in that same language (Italian input → Italian output, English input → English output, etc.). "
        "Output ONLY valid JSON — no markdown, no extra text."
    )
    user_prompt = f"""Context: {context_line}

Competitor Reviews:
{req.reviews[:4000]}

Analyze these reviews and generate EXACTLY 2 distinct customer avatars. Return JSON:
{{
  "avatars": [
    {{
      "name": "First name + last initial (fictional, e.g. Sofia M.)",
      "age_range": "e.g. 28-35",
      "occupation": "job or life stage",
      "tagline": "one sentence who they are",
      "psychographics": ["3-4 key personality traits/values"],
      "main_goal": "the #1 thing they want from this book",
      "pain_points": ["3 specific frustrations/struggles mentioned in reviews"],
      "buying_triggers": ["2-3 things that made them click Buy"],
      "language_patterns": ["3-4 exact phrases/words they use — quote style from reviews"],
      "content_implication": "what chapter topics/tone will resonate most with this avatar"
    }},
    {{ ...second avatar... }}
  ],
  "common_themes": ["top 3 themes across ALL reviews"],
  "positioning_insight": "one key differentiation opportunity based on what reviews LACK or complain about"
}}"""

    try:
        msg = claude.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=2000,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
        )
        raw = msg.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
        data = json.loads(raw)
        return data
    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="Risposta Claude non valida — riprova")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ══════════════════════════════════════════════════════════════
# MODI PER VINCERE — NICHE OPPORTUNITY FRAMEWORK
# ══════════════════════════════════════════════════════════════

class NicheOpportunityRequest(BaseModel):
    niche: str
    book_title: Optional[str] = ""
    book_type: Optional[str] = ""
    competitor_titles: Optional[str] = ""

@app.post("/api/niche-opportunity", dependencies=[_AUTH])
async def niche_opportunity(req: NicheOpportunityRequest):
    """Evaluate niche opportunity using the 3-question 'Modi per Vincere' framework."""
    context = f"Niche: {req.niche}"
    if req.book_title:
        context += f" | Book title idea: {req.book_title}"
    if req.book_type:
        context += f" | Book type: {req.book_type}"

    system_prompt = (
        "You are an Amazon KDP market strategist. Apply the 'Ways to Win' framework to evaluate "
        "a niche opportunity. Be specific, actionable, and honest — avoid generic advice. "
        "IMPORTANT: Detect the language of the niche/title/competitor titles provided, then write "
        "ALL text fields in that same language (Italian input → Italian output, English → English, etc.). "
        "Output ONLY valid JSON."
    )
    competitor_section = ""
    if req.competitor_titles:
        competitor_section = f"\nExisting competitor titles:\n{req.competitor_titles[:1500]}"

    user_prompt = f"""Evaluate this KDP niche opportunity: {context}{competitor_section}

Apply the 3-question framework and return JSON:
{{
  "gap_analysis": {{
    "has_gap": true/false,
    "gap_description": "specific underserved angle or format gap in this niche",
    "evidence": "why this gap exists based on what's missing or over-represented",
    "opportunity_score": 1-10
  }},
  "niche_down": {{
    "opportunities": [
      {{"angle": "niche-down idea", "example_title": "example book title", "target": "who exactly"}},
      {{"angle": "...", "example_title": "...", "target": "..."}},
      {{"angle": "...", "example_title": "...", "target": "..."}}
    ],
    "best_pick": "which of the 3 is strongest and why"
  }},
  "usp_angles": [
    {{"angle": "unique selling angle", "positioning_line": "one-line pitch", "differentiator": "what makes it stand out"}},
    {{"angle": "...", "positioning_line": "...", "differentiator": "..."}},
    {{"angle": "...", "positioning_line": "...", "differentiator": "..."}}
  ],
  "verdict": "overall 2-sentence assessment: should they enter this niche, and how?",
  "red_flags": ["any warning signs — saturation, declining trend, etc."]
}}"""

    try:
        msg = claude.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=2000,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
        )
        raw = msg.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
        data = json.loads(raw)
        return data
    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="Risposta Claude non valida — riprova")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ══════════════════════════════════════════════════════════════
# COMPETITION MAP — top 5 collective gap analysis
# ══════════════════════════════════════════════════════════════

@app.post("/api/competition-map", dependencies=[_AUTH])
async def competition_map(req: dict):
    niche = (req.get("niche") or "").strip()
    marketplace = req.get("marketplace", "us")
    if not niche:
        raise HTTPException(status_code=400, detail="Niche required")

    tld_map = {"us": "com", "de": "de", "it": "it", "es": "es", "fr": "fr", "uk": "co.uk"}
    tld = tld_map.get(marketplace, "com")

    books_data: list[dict] = []
    apify_used = False

    if APIFY_TOKEN:
        from urllib.parse import quote as url_quote
        search_url = f"https://www.amazon.{tld}/s?k={url_quote(niche)}&rh=n%3A283155"
        try:
            items = await asyncio.wait_for(
                run_actor("junglee/amazon-crawler", {
                    "categoryOrProductUrls": [{"url": search_url}],
                    "maxItemsPerStartUrl": 10,
                    "proxyConfiguration": {"useApifyProxy": True},
                }),
                timeout=90.0,
            )
            for item in (items or []):
                title = item.get("title") or item.get("name") or ""
                reviews = item.get("reviewsCount") or item.get("numberOfReviews") or 0
                price = item.get("price") or item.get("price_string") or ""
                bsr_raw = (item.get("bestsellersRank") or item.get("bsr")
                           or item.get("bestSellersRank") or [])
                bsr_val = None
                if isinstance(bsr_raw, list) and bsr_raw:
                    first = bsr_raw[0]
                    bsr_val = first.get("rank") or first.get("position") if isinstance(first, dict) else first
                elif isinstance(bsr_raw, int):
                    bsr_val = bsr_raw
                if title:
                    books_data.append({"title": title, "reviews": reviews,
                                       "bsr": bsr_val, "price": str(price)})
            apify_used = bool(books_data)
        except Exception:
            pass

    books_list = "\n".join(
        f"- \"{b['title']}\" (reviews: {b.get('reviews','?')}, BSR: {b.get('bsr','?')})"
        for b in books_data[:8]
    ) if books_data else "No live data — use your knowledge of this niche."

    prompt = f"""You are an Amazon KDP competitive intelligence expert.
Niche: {niche} | Marketplace: amazon.{tld}

Top books in this niche:
{books_list}

Analyze the competitive landscape and return JSON:
{{
  "books": [
    {{
      "title": "book title (use real titles if data available, otherwise estimate)",
      "reviews_est": "number or range",
      "price_est": "$X.XX or range",
      "weakness": "specific weakness or gap THIS book has",
      "reader_complaint": "what readers actually complain about"
    }}
  ],
  "collective_gap": "What ALL these books fail to do — the shared blind spot that no one has solved",
  "white_space_angle": "The specific book angle that doesn't exist yet but readers clearly want",
  "attack_strategy": "Concrete 2-sentence strategy to win against these books",
  "best_title_formula": "Title pattern that would dominate: example + explanation of the formula"
}}

Return exactly 5 books. Detect language from the niche and write all text in that language.
Output ONLY valid JSON."""

    try:
        raw = await call_claude(prompt, max_tokens=2200, allow_truncated=True)
        data = parse_json_safe(raw)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"competition-map Claude error: {e}")
    data["apify_used"] = apify_used
    return data


# ══════════════════════════════════════════════════════════════
# KEYWORD FUNNEL — head / body / long-tail tiers
# ══════════════════════════════════════════════════════════════

@app.post("/api/keyword-funnel", dependencies=[_AUTH])
async def keyword_funnel(req: dict):
    niche = (req.get("niche") or "").strip()
    marketplace = req.get("marketplace", "us")
    if not niche:
        raise HTTPException(status_code=400, detail="Niche required")

    lang_map = {"us": "en", "de": "de", "it": "it", "es": "es", "fr": "fr"}
    lang = lang_map.get(marketplace, "en")

    all_suggestions: set[str] = set()
    queries = [niche, f"{niche} book", f"{niche} for", f"{niche} guide", f"best {niche}"]

    async def fetch_one(q: str) -> list[str]:
        try:
            return await asyncio.wait_for(fetch_google_autocomplete(q, lang), timeout=8.0)
        except Exception:
            return []

    results = await asyncio.gather(*[fetch_one(q) for q in queries], return_exceptions=True)
    for r in results:
        if isinstance(r, list):
            all_suggestions.update(r)

    try:
        amazon_sugg = await asyncio.wait_for(fetch_amazon_autocomplete(niche), timeout=8.0)
        all_suggestions.update(amazon_sugg)
    except Exception:
        pass

    sugg_text = "\n".join(f"- {s}" for s in sorted(all_suggestions)[:50])

    prompt = f"""You are an Amazon KDP keyword strategy expert.
Niche: {niche} | Marketplace: {marketplace}

Autocomplete suggestions gathered:
{sugg_text if sugg_text else 'No live suggestions — use your knowledge.'}

Classify and expand keywords, return JSON:
{{
  "head": [
    {{"kw": "keyword", "difficulty": "High", "why": "why hard to rank", "monthly_est": "rough volume"}}
  ],
  "body": [
    {{"kw": "keyword", "difficulty": "Medium", "why": "why achievable", "monthly_est": "..."}}
  ],
  "longtail": [
    {{"kw": "keyword", "difficulty": "Low", "why": "easy win reason", "monthly_est": "..."}}
  ],
  "title_formula": "Recommended title using a body keyword — show a real example",
  "subtitle_formula": "Recommended subtitle weaving 1-2 longtail keywords — real example",
  "backend_keywords": ["kw1", "kw2", "kw3", "kw4", "kw5", "kw6", "kw7"],
  "insight": "One key strategic insight about keywords in this niche"
}}

Rules:
- head: 2-3 items, 1-2 word terms, everyone targets them
- body: 3-5 items, 2-3 word phrases, your title/subtitle sweet spot
- longtail: 4-6 items, 3-5 word phrases, easy-ranking gems
- backend_keywords: exactly 7 items for KDP backend keyword boxes
- ALL text in the language matching the niche input
Output ONLY valid JSON."""

    raw = await call_claude(prompt, max_tokens=1000)
    return parse_json_safe(raw)


# ══════════════════════════════════════════════════════════════
# ENTRY POINT
# ══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    import uvicorn

    env_file = os.path.join(os.path.dirname(__file__), ".env")
    if os.path.exists(env_file):
        print(f"[KDP Studio] Loading .env")
        with open(env_file) as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    key, _, val = line.partition("=")
                    os.environ[key.strip()] = val.strip().strip('"').strip("'")
        import sys
        _m = sys.modules[__name__]
        _m.ANTHROPIC_KEY     = env("ANTHROPIC_API_KEY")
        _m.REDDIT_USER_AGENT = env("REDDIT_USER_AGENT", "KDPStudio/1.0 (personal use, no auth)")
        _m.claude = anthropic.Anthropic(api_key=_m.ANTHROPIC_KEY)

    print("\n🚀 KDP Studio Backend v2")
    print(f"   Claude API:      {'OK' if ANTHROPIC_KEY else 'MISSING ANTHROPIC_API_KEY'}")
    print(f"   Apify:           {'OK' if APIFY_TOKEN else 'MISSING APIFY_TOKEN (Apify features disabled)'}")
    print(f"   Reddit mode:     Public JSON (no API key needed)")
    print(f"   Seed pool:       {sum(len(v) for v in NICHE_SEEDS.values())} seeds across {len(NICHE_SEEDS)} niches")
    print(f"   Subreddit pool:  {sum(len(v) for v in NICHE_SUBREDDIT_POOL.values())} subreddits across {len(NICHE_SUBREDDIT_POOL)} niches")
    print(f"   Opening styles:  {len(OPENING_STYLES)} rotating")
    print(f"   Tone variants:   {len(TONE_DESCRIPTORS)} rotating")
    print(f"\n   Docs:        http://localhost:8000/docs")
    print(f"   Zero-bias:   http://localhost:8000/discover\n")
    print(f"   Health: http://localhost:8000/health\n")

    port = int(os.environ.get("PORT", 8000))
    print(f"   Port: {port}")
    uvicorn.run(app, host="0.0.0.0", port=port, reload=False)
